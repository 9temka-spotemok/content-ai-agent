"""
Streamlit веб-интерфейс для Content AI Agent
"""
import streamlit as st
from datetime import datetime, timedelta
from typing import Dict, Any, List
import json
import io
import os

from content_ai_agent.main import ContentAIAgent
from content_ai_agent.modules.content_agent import ContentStatus, ContentFormat
from content_ai_agent.modules.analytics import AnalyticsReport
from content_ai_agent.config import Config

# Импорты для генерации документов
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Функция для создания OpenAI клиента
def create_openai_client(api_key: str = None):
    """Создает OpenAI клиент, если API ключ доступен, и проверяет его работоспособность"""
    # Используем переданный ключ, затем проверяем переменную окружения динамически
    key_to_use = api_key or Config.get_openai_api_key() or os.getenv("OPENAI_API_KEY", "").strip()
    
    if not key_to_use:
        return None
    
    # Проверяем формат ключа
    key_clean = key_to_use.strip()
    if not key_clean.startswith("sk-proj-") and not key_clean.startswith("sk-"):
        st.error("❌ Неверный формат API ключа. Ключ должен начинаться с 'sk-proj-' или 'sk-'")
        return None
    
    if len(key_clean) < 50:
        st.warning("⚠️ API ключ слишком короткий. Возможно, ключ обрезан при копировании.")
    
    try:
        from openai import OpenAI
    except ImportError:
        st.warning("Библиотека openai не установлена. Установите: pip install openai")
        return None
    
    client = OpenAI(api_key=key_clean)
    
    # Ленивая проверка: пробуем сделать простой запрос один раз за сессию
    if "openai_checked" not in st.session_state:
        st.session_state.openai_checked = False
        st.session_state.openai_ok = False
    
    if not st.session_state.openai_checked or api_key:  # Если передан новый ключ, проверяем заново
        try:
            # Лёгкий запрос: получаем список моделей (без передачи пользовательских данных)
            client.models.list()
            st.session_state.openai_ok = True
            # Если ключ был передан и работает, обновляем конфиг
            if api_key:
                import os
                os.environ["OPENAI_API_KEY"] = key_clean
                # Обновляем конфиг (используем глобальный импорт)
                Config.OPENAI_API_KEY = key_clean
        except Exception as e:
            st.session_state.openai_ok = False
            error_str = str(e)
            # Краткое сообщение об ошибке без подробного блока
            st.error("❌ Не удалось подключиться к OpenAI API. Проверьте корректность ключа и наличие сети.")
            # (техническую деталь логируем только в консоль, чтобы не засорять UI)
            try:
                print(f"[OpenAI error] {error_str}")
            except Exception:
                pass
            client = None
        finally:
            st.session_state.openai_checked = True
    
    return client

# Вспомогательные функции для цветных сообщений
def show_success_message(message: str):
    """Показывает сообщение об успехе с зеленым фоном"""
    st.markdown(
        f"""
        <div style="background-color: #d4edda; color: #155724; padding: 1rem; border-radius: 4px; margin: 1rem 0; border: 1px solid #c3e6cb;">
            {message}
        </div>
        """,
        unsafe_allow_html=True
    )

def show_error_message(message: str):
    """Показывает сообщение об ошибке с красным фоном"""
    st.markdown(
        f"""
        <div style="background-color: #f8d7da; color: #721c24; padding: 1rem; border-radius: 4px; margin: 1rem 0; border: 1px solid #f5c6cb;">
            {message}
        </div>
        """,
        unsafe_allow_html=True
    )
    
def show_ai_assistant(message: str, theme: str = "light"):
    """Показывает AI-ассистента с облаком текста (без фото и имени человека)"""
    text_color = "#000000" if theme == "light" else "#ffffff"
    bg_color = "#ffffff" if theme == "light" else "#1a1a1a"
    
    # Экранируем сообщение для HTML
    import html
    message_escaped = html.escape(message)
    
    # Стили для облака с текстом
    css_style = f"""
    <style>
    @keyframes fadeIn {{
        from {{ opacity: 0; transform: translateY(-10px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    @keyframes bubblePulse {{
        0%, 100% {{ transform: scale(1); box-shadow: 0 4px 8px rgba(0,0,0,0.1); }}
        50% {{ transform: scale(1.02); box-shadow: 0 6px 12px rgba(0,0,0,0.15); }}
    }}
    .ai-assistant-wrapper {{
        display: flex;
        flex-direction: column;
        align-items: center;
        margin: 2rem 0;
        animation: fadeIn 0.6s ease-in;
    }}
    .speech-bubble-wrapper {{
        position: relative;
        background-color: {bg_color};
        border: 2px solid {text_color};
        border-radius: 20px;
        padding: 1.2rem 1.8rem;
        margin-bottom: 1.5rem;
        max-width: 550px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
        animation: fadeIn 0.6s ease-in, bubblePulse 3s ease-in-out infinite;
    }}
    .speech-bubble-wrapper::after {{
        content: '';
        position: absolute;
        bottom: -18px;
        left: 50%;
        transform: translateX(-50%);
        width: 0;
        height: 0;
        border-left: 18px solid transparent;
        border-right: 18px solid transparent;
        border-top: 18px solid {bg_color};
        z-index: 2;
    }}
    .speech-bubble-wrapper::before {{
        content: '';
        position: absolute;
        bottom: -20px;
        left: 50%;
        transform: translateX(-50%);
        width: 0;
        height: 0;
        border-left: 20px solid transparent;
        border-right: 20px solid transparent;
        border-top: 20px solid {text_color};
        z-index: 1;
    }}
    .speech-text-wrapper {{
        color: {text_color};
        font-size: 1.05rem;
        line-height: 1.7;
        margin: 0;
        text-align: center;
        font-weight: 400;
    }}
    </style>
    """
    
    # Выводим стили
    st.markdown(css_style, unsafe_allow_html=True)
    
    # Облако с текстом
    st.markdown(
        f"""
        <div class="ai-assistant-wrapper">
            <div class="speech-bubble-wrapper">
                <p class="speech-text-wrapper">{message_escaped}</p>
        </div>
        """,
        unsafe_allow_html=True
    )

# Функции для генерации документов
def generate_docx_report(hypothesis_levels: Dict, product_info: Dict = None) -> bytes:
    """Генерирует DOCX отчет с выбранными гипотезами"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчет по анализу целевой аудитории', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о продукте
    if product_info:
        doc.add_heading('Информация о продукте', level=1)
        doc.add_paragraph(f"Название: {product_info.get('name', 'Не указано')}")
        doc.add_paragraph(f"Категория: {product_info.get('category', 'Не указано')}")
        doc.add_paragraph(f"Описание: {product_info.get('description', 'Не указано')}")
        doc.add_paragraph('')
    
    # Выбранные гипотезы
    doc.add_heading('Выбранные гипотезы', level=1)
    
    level_names = {
        "market": "Рынок",
        "niche": "Ниша",
        "audience": "Аудитория",
        "segment": "Сегмент",
        "pain": "Боль",
        "task": "Задача"
    }
    
    levels_order = ["market", "niche", "audience", "segment", "pain", "task"]
    
    for level in levels_order:
        level_data = hypothesis_levels.get(level, {})
        if level_data.get("completed") and level_data.get("selected"):
            selected = level_data["selected"]
            doc.add_heading(level_names[level], level=2)
            
            name = selected.get("name", "Не указано")
            description = selected.get("description", "")
            characteristics = selected.get("characteristics", [])
            potential = selected.get("potential", "")
            
            doc.add_paragraph(f"Название: {name}", style='List Bullet')
            if description:
                doc.add_paragraph(f"Описание: {description}")
            if characteristics:
                doc.add_paragraph("Характеристики:")
                for char in characteristics:
                    doc.add_paragraph(char, style='List Bullet 2')
            if potential:
                doc.add_paragraph(f"Потенциал: {potential}")
            doc.add_paragraph('')
    
    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_pdf_report(hypothesis_levels: Dict, product_info: Dict = None) -> bytes:
    """Генерирует PDF отчет с выбранными гипотезами"""
    if not PDF_AVAILABLE:
        raise ImportError("reportlab не установлен. Установите: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Создаем кастомные стили
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor='black',
        spaceAfter=24,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='black',
        spaceAfter=8,
        spaceBefore=16,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=6,
        leading=14,
        leftIndent=0,
        rightIndent=0
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=4,
        leading=14,
        leftIndent=20,
        bulletIndent=10
    )
    
    # Функция для экранирования HTML
    def escape_html(text):
        if not text:
            return ""
        return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # Заголовок
    story.append(Paragraph('Отчет по анализу целевой аудитории', title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Информация о продукте
    if product_info:
        story.append(Paragraph('Информация о продукте', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        name = escape_html(product_info.get('name', 'Не указано'))
        category = escape_html(product_info.get('category', 'Не указано'))
        description = escape_html(product_info.get('description', 'Не указано'))
        
        story.append(Paragraph(f"<b>Название:</b> {name}", normal_style))
        story.append(Paragraph(f"<b>Категория:</b> {category}", normal_style))
        if description:
            story.append(Paragraph(f"<b>Описание:</b> {description}", normal_style))
        story.append(Spacer(1, 0.2*inch))
    
    # Выбранные гипотезы
    story.append(Paragraph('Выбранные гипотезы', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    
    level_names = {
        "market": "1. Рынок",
        "niche": "2. Ниша",
        "audience": "3. Аудитория",
        "segment": "4. Сегмент",
        "pain": "5. Боль",
        "task": "6. Задача"
    }
    
    levels_order = ["market", "niche", "audience", "segment", "pain", "task"]
    
    for level in levels_order:
        level_data = hypothesis_levels.get(level, {})
        if level_data.get("completed") and level_data.get("selected"):
            selected = level_data["selected"]
            story.append(Paragraph(level_names[level], heading2_style))
            
            name = escape_html(selected.get("name", "Не указано"))
            description = escape_html(selected.get("description", ""))
            characteristics = selected.get("characteristics", [])
            potential = escape_html(selected.get("potential", ""))
            
            story.append(Paragraph(f"<b>Название:</b> {name}", normal_style))
            story.append(Spacer(1, 0.1*inch))
            
            if description:
                story.append(Paragraph(f"<b>Описание:</b>", normal_style))
                story.append(Paragraph(description, normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if characteristics:
                story.append(Paragraph("<b>Характеристики:</b>", normal_style))
                for char in characteristics:
                    char_escaped = escape_html(char)
                    story.append(Paragraph(f"• {char_escaped}", bullet_style))
                story.append(Spacer(1, 0.1*inch))
            
            if potential:
                story.append(Paragraph(f"<b>Потенциал:</b> {potential}", normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.15*inch))
    
    # Дата создания
    story.append(Spacer(1, 0.3*inch))
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor='#666666',
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}", date_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_analytics_docx_report(
    report: AnalyticsReport, 
    days: int,
    product_info: Dict = None,
    recommendations: List[str] = None
) -> bytes:
    """Генерирует DOCX отчет по аналитике"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчет по аналитике контента', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о периоде
    doc.add_paragraph(f"Период анализа: {days} дней")
    doc.add_paragraph(f"С {report.period_start.strftime('%d.%m.%Y')} по {report.period_end.strftime('%d.%m.%Y')}")
    doc.add_paragraph('')
    
    # Информация о продукте
    if product_info:
        doc.add_heading('Информация о продукте', level=1)
        doc.add_paragraph(f"Название: {product_info.get('name', 'Не указано')}")
        doc.add_paragraph(f"Категория: {product_info.get('category', 'Не указано')}")
        if product_info.get('description'):
            doc.add_paragraph(f"Описание: {product_info.get('description', 'Не указано')}")
        doc.add_paragraph('')
    
    # Основные метрики
    doc.add_heading('Основные метрики', level=1)
    doc.add_paragraph(f"Всего постов: {report.total_posts}")
    doc.add_paragraph(f"Просмотры: {report.total_views:,}")
    doc.add_paragraph(f"Вовлеченность: {report.total_engagement:,}")
    doc.add_paragraph(f"Engagement Rate: {report.avg_engagement_rate:.2f}%")
    doc.add_paragraph('')
    
    # Топ контент
    if report.top_performing_content:
        doc.add_heading('Топ контент', level=1)
        for idx, item in enumerate(report.top_performing_content, 1):
            doc.add_heading(f"{idx}. {item.get('title', 'Без названия')}", level=2)
            doc.add_paragraph(f"Engagement: {item.get('engagement', 0):,}")
            doc.add_paragraph(f"Engagement Rate: {item.get('engagement_rate', 0):.2f}%")
            doc.add_paragraph('')
    
    # Производительность по каналам
    if report.channel_performance:
        doc.add_heading('Производительность по каналам', level=1)
        for channel, perf in report.channel_performance.items():
            doc.add_heading(channel, level=2)
            doc.add_paragraph(f"Посты: {perf.get('posts', 0)}")
            doc.add_paragraph(f"Просмотры: {perf.get('views', 0):,}")
            doc.add_paragraph(f"Вовлеченность: {perf.get('engagement', 0):,}")
            doc.add_paragraph(f"Engagement Rate: {perf.get('avg_engagement_rate', 0):.2f}%")
            doc.add_paragraph('')
    
    # Производительность по столпам
    if report.pillar_performance:
        doc.add_heading('Производительность по столпам контента', level=1)
        for pillar, perf in report.pillar_performance.items():
            doc.add_paragraph(f"{pillar}: {perf.get('posts', 0)} постов")
        doc.add_paragraph('')
    
    # Производительность по стадиям воронки
    if report.funnel_stage_performance:
        doc.add_heading('Производительность по стадиям воронки', level=1)
        for stage, perf in report.funnel_stage_performance.items():
            doc.add_paragraph(f"{stage}: {perf.get('posts', 0)} постов")
        doc.add_paragraph('')
    
    # Рекомендации
    if recommendations:
        doc.add_heading('Рекомендации', level=1)
        for idx, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{idx}. {rec}", style='List Bullet')
        doc.add_paragraph('')
    
    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_analytics_pdf_report(
    report: AnalyticsReport,
    days: int,
    product_info: Dict = None,
    recommendations: List[str] = None
) -> bytes:
    """Генерирует PDF отчет по аналитике"""
    if not PDF_AVAILABLE:
        raise ImportError("reportlab не установлен. Установите: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Создаем кастомные стили
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor='black',
        spaceAfter=24,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='black',
        spaceAfter=8,
        spaceBefore=16,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=6,
        leading=14,
        leftIndent=0,
        rightIndent=0
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=4,
        leading=14,
        leftIndent=20,
        bulletIndent=10
    )
    
    # Функция для экранирования HTML
    def escape_html(text):
        if not text:
            return ""
        return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # Заголовок
    story.append(Paragraph('Отчет по аналитике контента', title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Информация о периоде
    period_text = f"Период анализа: {days} дней<br/>С {report.period_start.strftime('%d.%m.%Y')} по {report.period_end.strftime('%d.%m.%Y')}"
    story.append(Paragraph(period_text, normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Информация о продукте
    if product_info:
        story.append(Paragraph('Информация о продукте', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        name = escape_html(product_info.get('name', 'Не указано'))
        category = escape_html(product_info.get('category', 'Не указано'))
        description = escape_html(product_info.get('description', ''))
        
        story.append(Paragraph(f"<b>Название:</b> {name}", normal_style))
        story.append(Paragraph(f"<b>Категория:</b> {category}", normal_style))
        if description:
            story.append(Paragraph(f"<b>Описание:</b> {description}", normal_style))
        story.append(Spacer(1, 0.2*inch))
    
    # Основные метрики
    story.append(Paragraph('Основные метрики', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    
    story.append(Paragraph(f"<b>Всего постов:</b> {report.total_posts}", normal_style))
    story.append(Paragraph(f"<b>Просмотры:</b> {report.total_views:,}", normal_style))
    story.append(Paragraph(f"<b>Вовлеченность:</b> {report.total_engagement:,}", normal_style))
    story.append(Paragraph(f"<b>Engagement Rate:</b> {report.avg_engagement_rate:.2f}%", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Топ контент
    if report.top_performing_content:
        story.append(Paragraph('Топ контент', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        for idx, item in enumerate(report.top_performing_content, 1):
            title = escape_html(item.get('title', 'Без названия'))
            engagement = item.get('engagement', 0)
            engagement_rate = item.get('engagement_rate', 0)
            
            story.append(Paragraph(f"{idx}. {title}", heading2_style))
            story.append(Paragraph(f"<b>Engagement:</b> {engagement:,}", normal_style))
            story.append(Paragraph(f"<b>Engagement Rate:</b> {engagement_rate:.2f}%", normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.1*inch))
    
    # Производительность по каналам
    if report.channel_performance:
        story.append(Paragraph('Производительность по каналам', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        for channel, perf in report.channel_performance.items():
            channel_escaped = escape_html(channel)
            story.append(Paragraph(channel_escaped, heading2_style))
            story.append(Paragraph(f"<b>Посты:</b> {perf.get('posts', 0)}", normal_style))
            story.append(Paragraph(f"<b>Просмотры:</b> {perf.get('views', 0):,}", normal_style))
            story.append(Paragraph(f"<b>Вовлеченность:</b> {perf.get('engagement', 0):,}", normal_style))
            story.append(Paragraph(f"<b>Engagement Rate:</b> {perf.get('avg_engagement_rate', 0):.2f}%", normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.1*inch))
    
    # Производительность по столпам
    if report.pillar_performance:
        story.append(Paragraph('Производительность по столпам контента', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        for pillar, perf in report.pillar_performance.items():
            pillar_escaped = escape_html(pillar)
            story.append(Paragraph(f"<b>{pillar_escaped}:</b> {perf.get('posts', 0)} постов", normal_style))
        
        story.append(Spacer(1, 0.2*inch))
    
    # Производительность по стадиям воронки
    if report.funnel_stage_performance:
        story.append(Paragraph('Производительность по стадиям воронки', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        for stage, perf in report.funnel_stage_performance.items():
            stage_escaped = escape_html(stage)
            story.append(Paragraph(f"<b>{stage_escaped}:</b> {perf.get('posts', 0)} постов", normal_style))
        
        story.append(Spacer(1, 0.2*inch))
    
    # Рекомендации
    if recommendations:
        story.append(Paragraph('Рекомендации', heading1_style))
        story.append(Spacer(1, 0.15*inch))
        
        for idx, rec in enumerate(recommendations, 1):
            rec_escaped = escape_html(rec)
            story.append(Paragraph(f"{idx}. {rec_escaped}", bullet_style))
        
        story.append(Spacer(1, 0.2*inch))
    
    # Дата создания
    story.append(Spacer(1, 0.3*inch))
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor='#666666',
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}", date_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_comprehensive_docx_report(
    product_info: Dict = None,
    deep_impact_audit: Dict = None,
    hypothesis_levels: Dict = None,
    strategy = None,
    content_list: List = None,
    scheduling_data: Dict = None,
    schedule: Dict = None,
    analytics_report = None,
    analytics_days: int = 30,
    recommendations: List[str] = None
) -> bytes:
    """Генерирует комплексный DOCX отчет по всем разделам"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Комплексный отчет по контент-маркетингу', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph('')
    
    # 1. ОНБОРДИНГ
    doc.add_heading('1. Онбординг', level=1)
    if product_info:
        doc.add_paragraph(f"Название продукта: {product_info.get('name', 'Не указано')}")
        doc.add_paragraph(f"Категория: {product_info.get('category', 'Не указано')}")
        if product_info.get('description'):
            doc.add_paragraph(f"Описание: {product_info.get('description', 'Не указано')}")
        if product_info.get('features'):
            doc.add_paragraph("Функции:")
            for feature in product_info.get('features', []):
                doc.add_paragraph(f"  • {feature}", style='List Bullet 2')
        if product_info.get('target_audience'):
            doc.add_paragraph(f"Целевая аудитория: {product_info.get('target_audience', 'Не указано')}")
    else:
        doc.add_paragraph("Информация о продукте не заполнена")
    doc.add_paragraph('')
    
    # 2. DEEP IMPACT
    doc.add_heading('2. DEEP IMPACT', level=1)
    if deep_impact_audit:
        doc.add_paragraph(f"Бизнес-ниша: {deep_impact_audit.get('business_niche', 'Не указано')}")
        doc.add_paragraph(f"Стадия бизнеса: {deep_impact_audit.get('business_stage', 'Не указано')}")
        if deep_impact_audit.get('existing_data'):
            doc.add_paragraph(f"Существующие данные: {deep_impact_audit.get('existing_data', 'Не указано')}")
    else:
        doc.add_paragraph("Аудит DEEP IMPACT не проведен")
    doc.add_paragraph('')
    
    # Гипотезы
    if hypothesis_levels:
        doc.add_heading('2.1. Выбранные гипотезы', level=2)
        level_names = {
            "market": "Рынок",
            "niche": "Ниша",
            "audience": "Аудитория",
            "segment": "Сегмент",
            "pain": "Боль",
            "task": "Задача"
        }
        levels_order = ["market", "niche", "audience", "segment", "pain", "task"]
        
        for level in levels_order:
            level_data = hypothesis_levels.get(level, {})
            if level_data.get("completed") and level_data.get("selected"):
                selected = level_data["selected"]
                doc.add_heading(level_names[level], level=3)
                
                name = selected.get("name", "Не указано")
                description = selected.get("description", "")
                characteristics = selected.get("characteristics", [])
                potential = selected.get("potential", "")
                
                doc.add_paragraph(f"Название: {name}", style='List Bullet')
                if description:
                    doc.add_paragraph(f"Описание: {description}")
                if characteristics:
                    doc.add_paragraph("Характеристики:")
                    for char in characteristics:
                        doc.add_paragraph(char, style='List Bullet 2')
                if potential:
                    doc.add_paragraph(f"Потенциал: {potential}")
                doc.add_paragraph('')
    else:
        doc.add_paragraph("Гипотезы не сгенерированы")
    doc.add_paragraph('')
    
    # 3. СТРАТЕГИЯ
    doc.add_heading('3. Контент-стратегия', level=1)
    if strategy:
        doc.add_heading('3.1. Целевая аудитория', level=2)
        if hasattr(strategy, 'target_audience') and strategy.target_audience:
            for key, value in strategy.target_audience.items():
                doc.add_paragraph(f"{key}: {value}")
        
        doc.add_heading('3.2. Контентные столпы', level=2)
        if hasattr(strategy, 'content_pillars') and strategy.content_pillars:
            for idx, pillar in enumerate(strategy.content_pillars, 1):
                if isinstance(pillar, str):
                    doc.add_paragraph(f"{idx}. {pillar}")
                else:
                    doc.add_paragraph(f"{idx}. {pillar.name if hasattr(pillar, 'name') else str(pillar)}")
        
        doc.add_heading('3.3. Каналы', level=2)
        if hasattr(strategy, 'channels') and strategy.channels:
            for channel in strategy.channels:
                doc.add_paragraph(f"• {channel}")
        
        doc.add_heading('3.4. Tone of Voice', level=2)
        if hasattr(strategy, 'tone_of_voice'):
            doc.add_paragraph(strategy.tone_of_voice)
        
        doc.add_heading('3.5. Ключевые сообщения', level=2)
        if hasattr(strategy, 'key_messages') and strategy.key_messages:
            for msg in strategy.key_messages:
                doc.add_paragraph(f"• {msg}")
        
        doc.add_heading('3.6. Воронка контента', level=2)
        if hasattr(strategy, 'content_funnel') and strategy.content_funnel:
            for stage, topics in strategy.content_funnel.items():
                doc.add_paragraph(f"{stage}:")
                for topic in topics:
                    doc.add_paragraph(f"  • {topic}", style='List Bullet 2')
    else:
        doc.add_paragraph("Контент-стратегия не создана")
    doc.add_paragraph('')
    
    # 4. КОНТЕНТ
    doc.add_heading('4. Сгенерированный контент', level=1)
    if content_list:
        doc.add_paragraph(f"Всего постов: {len(content_list)}")
        doc.add_paragraph('')
        for idx, content in enumerate(content_list, 1):
            doc.add_heading(f"{idx}. {content.title}", level=2)
            doc.add_paragraph(f"Формат: {content.format.value if hasattr(content.format, 'value') else content.format}")
            doc.add_paragraph(f"Канал: {content.channel}")
            doc.add_paragraph(f"Столп: {content.pillar}")
            doc.add_paragraph(f"Стадия воронки: {content.funnel_stage}")
            doc.add_paragraph(f"Статус: {content.status.value if hasattr(content.status, 'value') else content.status}")
            doc.add_paragraph("Содержание:")
            doc.add_paragraph(content.content)
            doc.add_paragraph('')
    else:
        doc.add_paragraph("Контент не сгенерирован")
    doc.add_paragraph('')
    
    # 5. ПЛАНИРОВАНИЕ
    doc.add_heading('5. Планирование публикаций', level=1)
    if scheduling_data:
        doc.add_paragraph(f"Каналы: {', '.join(scheduling_data.get('channels', []))}")
        if scheduling_data.get('start_date'):
            start_date = scheduling_data['start_date']
            if isinstance(start_date, str):
                doc.add_paragraph(f"Дата начала: {start_date}")
            else:
                doc.add_paragraph(f"Дата начала: {start_date.strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"Постов в неделю: {scheduling_data.get('posts_per_week', 'Не указано')}")
        doc.add_paragraph(f"Выбрано постов: {len(scheduling_data.get('selected_ids', []))}")
    else:
        doc.add_paragraph("Данные планирования не заполнены")
    doc.add_paragraph('')
    
    if schedule:
        doc.add_heading('5.1. Календарь публикаций', level=2)
        for channel, posts in schedule.items():
            doc.add_heading(f"{channel.upper()} - {len(posts)} постов", level=3)
            for post in posts[:20]:  # Показываем первые 20
                if hasattr(post, 'scheduled_time') and hasattr(post, 'content'):
                    time_str = post.scheduled_time.strftime('%Y-%m-%d %H:%M') if hasattr(post.scheduled_time, 'strftime') else str(post.scheduled_time)
                    title = post.content.title if hasattr(post.content, 'title') else str(post.content)
                    doc.add_paragraph(f"  • {time_str} - {title}", style='List Bullet 2')
            if len(posts) > 20:
                doc.add_paragraph(f"  ... и еще {len(posts) - 20} постов")
            doc.add_paragraph('')
    else:
        doc.add_paragraph("Расписание не создано")
    doc.add_paragraph('')
    
    # 6. АНАЛИТИКА
    doc.add_heading('6. Аналитика', level=1)
    if analytics_report:
        doc.add_paragraph(f"Период анализа: {analytics_days} дней")
        doc.add_paragraph(f"С {analytics_report.period_start.strftime('%d.%m.%Y')} по {analytics_report.period_end.strftime('%d.%m.%Y')}")
        doc.add_paragraph('')
        
        doc.add_heading('6.1. Основные метрики', level=2)
        doc.add_paragraph(f"Всего постов: {analytics_report.total_posts}")
        doc.add_paragraph(f"Просмотры: {analytics_report.total_views:,}")
        doc.add_paragraph(f"Вовлеченность: {analytics_report.total_engagement:,}")
        doc.add_paragraph(f"Engagement Rate: {analytics_report.avg_engagement_rate:.2f}%")
        doc.add_paragraph('')
        
        if analytics_report.top_performing_content:
            doc.add_heading('6.2. Топ контент', level=2)
            for idx, item in enumerate(analytics_report.top_performing_content, 1):
                doc.add_heading(f"{idx}. {item.get('title', 'Без названия')}", level=3)
                doc.add_paragraph(f"Engagement: {item.get('engagement', 0):,}")
                doc.add_paragraph(f"Engagement Rate: {item.get('engagement_rate', 0):.2f}%")
                doc.add_paragraph('')
        
        if analytics_report.channel_performance:
            doc.add_heading('6.3. Производительность по каналам', level=2)
            for channel, perf in analytics_report.channel_performance.items():
                doc.add_heading(channel, level=3)
                doc.add_paragraph(f"Посты: {perf.get('posts', 0)}")
                doc.add_paragraph(f"Просмотры: {perf.get('views', 0):,}")
                doc.add_paragraph(f"Вовлеченность: {perf.get('engagement', 0):,}")
                doc.add_paragraph(f"Engagement Rate: {perf.get('avg_engagement_rate', 0):.2f}%")
                doc.add_paragraph('')
        
        # Рекомендации
        if recommendations:
            doc.add_heading('6.4. Рекомендации', level=2)
            for idx, rec in enumerate(recommendations, 1):
                doc.add_paragraph(f"{idx}. {rec}", style='List Bullet')
    else:
        doc.add_paragraph("Аналитика не сгенерирована")
    doc.add_paragraph('')
    
    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_comprehensive_pdf_report(
    product_info: Dict = None,
    deep_impact_audit: Dict = None,
    hypothesis_levels: Dict = None,
    strategy = None,
    content_list: List = None,
    scheduling_data: Dict = None,
    schedule: Dict = None,
    analytics_report = None,
    analytics_days: int = 30,
    recommendations: List[str] = None
) -> bytes:
    """Генерирует комплексный PDF отчет по всем разделам"""
    if not PDF_AVAILABLE:
        raise ImportError("reportlab не установлен. Установите: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Создаем кастомные стили
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor='black',
        spaceAfter=24,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='black',
        spaceAfter=8,
        spaceBefore=16,
        fontName='Helvetica-Bold'
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor='black',
        spaceAfter=6,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=6,
        leading=14,
        leftIndent=0,
        rightIndent=0
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=4,
        leading=14,
        leftIndent=20,
        bulletIndent=10
    )
    
    # Функция для экранирования HTML
    def escape_html(text):
        if not text:
            return ""
        return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # Заголовок
    story.append(Paragraph('Комплексный отчет по контент-маркетингу', title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style))
    story.append(Spacer(1, 0.3*inch))
    
    # 1. ОНБОРДИНГ
    story.append(Paragraph('1. Онбординг', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if product_info:
        name = escape_html(product_info.get('name', 'Не указано'))
        category = escape_html(product_info.get('category', 'Не указано'))
        description = escape_html(product_info.get('description', ''))
        target_audience = escape_html(product_info.get('target_audience', ''))
        
        story.append(Paragraph(f"<b>Название продукта:</b> {name}", normal_style))
        story.append(Paragraph(f"<b>Категория:</b> {category}", normal_style))
        if description:
            story.append(Paragraph(f"<b>Описание:</b> {description}", normal_style))
        if product_info.get('features'):
            story.append(Paragraph("<b>Функции:</b>", normal_style))
            for feature in product_info.get('features', []):
                feature_escaped = escape_html(feature)
                story.append(Paragraph(f"• {feature_escaped}", bullet_style))
        if target_audience:
            story.append(Paragraph(f"<b>Целевая аудитория:</b> {target_audience}", normal_style))
    else:
        story.append(Paragraph("Информация о продукте не заполнена", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 2. DEEP IMPACT
    story.append(Paragraph('2. DEEP IMPACT', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if deep_impact_audit:
        niche = escape_html(deep_impact_audit.get('business_niche', 'Не указано'))
        stage = escape_html(deep_impact_audit.get('business_stage', 'Не указано'))
        existing_data = escape_html(deep_impact_audit.get('existing_data', ''))
        
        story.append(Paragraph(f"<b>Бизнес-ниша:</b> {niche}", normal_style))
        story.append(Paragraph(f"<b>Стадия бизнеса:</b> {stage}", normal_style))
        if existing_data:
            story.append(Paragraph(f"<b>Существующие данные:</b> {existing_data}", normal_style))
    else:
        story.append(Paragraph("Аудит DEEP IMPACT не проведен", normal_style))
    story.append(Spacer(1, 0.15*inch))
    
    # Гипотезы
    if hypothesis_levels:
        story.append(Paragraph('2.1. Выбранные гипотезы', heading2_style))
        story.append(Spacer(1, 0.1*inch))
        
        level_names = {
            "market": "Рынок",
            "niche": "Ниша",
            "audience": "Аудитория",
            "segment": "Сегмент",
            "pain": "Боль",
            "task": "Задача"
        }
        levels_order = ["market", "niche", "audience", "segment", "pain", "task"]
        
        for level in levels_order:
            level_data = hypothesis_levels.get(level, {})
            if level_data.get("completed") and level_data.get("selected"):
                selected = level_data["selected"]
                story.append(Paragraph(level_names[level], heading3_style))
                
                name = escape_html(selected.get("name", "Не указано"))
                description = escape_html(selected.get("description", ""))
                characteristics = selected.get("characteristics", [])
                potential = escape_html(selected.get("potential", ""))
                
                story.append(Paragraph(f"<b>Название:</b> {name}", normal_style))
                if description:
                    story.append(Paragraph(f"<b>Описание:</b> {description}", normal_style))
                if characteristics:
                    story.append(Paragraph("<b>Характеристики:</b>", normal_style))
                    for char in characteristics:
                        char_escaped = escape_html(char)
                        story.append(Paragraph(f"• {char_escaped}", bullet_style))
                if potential:
                    story.append(Paragraph(f"<b>Потенциал:</b> {potential}", normal_style))
                story.append(Spacer(1, 0.1*inch))
    else:
        story.append(Paragraph("Гипотезы не сгенерированы", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 3. СТРАТЕГИЯ
    story.append(Paragraph('3. Контент-стратегия', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if strategy:
        if hasattr(strategy, 'content_pillars') and strategy.content_pillars:
            story.append(Paragraph('3.1. Контентные столпы', heading2_style))
            for idx, pillar in enumerate(strategy.content_pillars, 1):
                pillar_str = pillar if isinstance(pillar, str) else (pillar.name if hasattr(pillar, 'name') else str(pillar))
                pillar_escaped = escape_html(pillar_str)
                story.append(Paragraph(f"{idx}. {pillar_escaped}", normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if hasattr(strategy, 'channels') and strategy.channels:
            story.append(Paragraph('3.2. Каналы', heading2_style))
            for channel in strategy.channels:
                channel_escaped = escape_html(channel)
                story.append(Paragraph(f"• {channel_escaped}", bullet_style))
            story.append(Spacer(1, 0.1*inch))
        
        if hasattr(strategy, 'tone_of_voice'):
            story.append(Paragraph('3.3. Tone of Voice', heading2_style))
            tone_escaped = escape_html(strategy.tone_of_voice)
            story.append(Paragraph(tone_escaped, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if hasattr(strategy, 'key_messages') and strategy.key_messages:
            story.append(Paragraph('3.4. Ключевые сообщения', heading2_style))
            for msg in strategy.key_messages:
                msg_escaped = escape_html(msg)
                story.append(Paragraph(f"• {msg_escaped}", bullet_style))
            story.append(Spacer(1, 0.1*inch))
    else:
        story.append(Paragraph("Контент-стратегия не создана", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 4. КОНТЕНТ
    story.append(Paragraph('4. Сгенерированный контент', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if content_list:
        story.append(Paragraph(f"<b>Всего постов:</b> {len(content_list)}", normal_style))
        story.append(Spacer(1, 0.1*inch))
        for idx, content in enumerate(content_list[:10], 1):  # Показываем первые 10
            title_escaped = escape_html(content.title)
            story.append(Paragraph(f"{idx}. {title_escaped}", heading3_style))
            story.append(Paragraph(f"<b>Формат:</b> {content.format.value if hasattr(content.format, 'value') else content.format}", normal_style))
            story.append(Paragraph(f"<b>Канал:</b> {content.channel}", normal_style))
            story.append(Paragraph(f"<b>Столп:</b> {content.pillar}", normal_style))
            content_text = escape_html(content.content[:500])  # Первые 500 символов
            story.append(Paragraph(f"<b>Содержание:</b> {content_text}...", normal_style))
            story.append(Spacer(1, 0.1*inch))
        if len(content_list) > 10:
            story.append(Paragraph(f"... и еще {len(content_list) - 10} постов", normal_style))
    else:
        story.append(Paragraph("Контент не сгенерирован", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 5. ПЛАНИРОВАНИЕ
    story.append(Paragraph('5. Планирование публикаций', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if scheduling_data:
        channels_str = ', '.join(scheduling_data.get('channels', []))
        story.append(Paragraph(f"<b>Каналы:</b> {channels_str}", normal_style))
        if scheduling_data.get('start_date'):
            start_date = scheduling_data['start_date']
            if isinstance(start_date, str):
                story.append(Paragraph(f"<b>Дата начала:</b> {start_date}", normal_style))
            else:
                story.append(Paragraph(f"<b>Дата начала:</b> {start_date.strftime('%d.%m.%Y')}", normal_style))
        story.append(Paragraph(f"<b>Постов в неделю:</b> {scheduling_data.get('posts_per_week', 'Не указано')}", normal_style))
        story.append(Paragraph(f"<b>Выбрано постов:</b> {len(scheduling_data.get('selected_ids', []))}", normal_style))
    else:
        story.append(Paragraph("Данные планирования не заполнены", normal_style))
    story.append(Spacer(1, 0.15*inch))
    
    if schedule:
        story.append(Paragraph('5.1. Календарь публикаций', heading2_style))
        for channel, posts in schedule.items():
            channel_escaped = escape_html(channel.upper())
            story.append(Paragraph(f"{channel_escaped} - {len(posts)} постов", heading3_style))
            for post in posts[:10]:  # Показываем первые 10
                if hasattr(post, 'scheduled_time') and hasattr(post, 'content'):
                    time_str = post.scheduled_time.strftime('%Y-%m-%d %H:%M') if hasattr(post.scheduled_time, 'strftime') else str(post.scheduled_time)
                    title = post.content.title if hasattr(post.content, 'title') else str(post.content)
                    title_escaped = escape_html(title)
                    story.append(Paragraph(f"• {time_str} - {title_escaped}", bullet_style))
            if len(posts) > 10:
                story.append(Paragraph(f"... и еще {len(posts) - 10} постов", normal_style))
            story.append(Spacer(1, 0.1*inch))
    else:
        story.append(Paragraph("Расписание не создано", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 6. АНАЛИТИКА
    story.append(Paragraph('6. Аналитика', heading1_style))
    story.append(Spacer(1, 0.15*inch))
    if analytics_report:
        story.append(Paragraph(f"<b>Период анализа:</b> {analytics_days} дней", normal_style))
        period_text = f"С {analytics_report.period_start.strftime('%d.%m.%Y')} по {analytics_report.period_end.strftime('%d.%m.%Y')}"
        story.append(Paragraph(period_text, normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph('6.1. Основные метрики', heading2_style))
        story.append(Paragraph(f"<b>Всего постов:</b> {analytics_report.total_posts}", normal_style))
        story.append(Paragraph(f"<b>Просмотры:</b> {analytics_report.total_views:,}", normal_style))
        story.append(Paragraph(f"<b>Вовлеченность:</b> {analytics_report.total_engagement:,}", normal_style))
        story.append(Paragraph(f"<b>Engagement Rate:</b> {analytics_report.avg_engagement_rate:.2f}%", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        if analytics_report.top_performing_content:
            story.append(Paragraph('6.2. Топ контент', heading2_style))
            for idx, item in enumerate(analytics_report.top_performing_content, 1):
                title_escaped = escape_html(item.get('title', 'Без названия'))
                story.append(Paragraph(f"{idx}. {title_escaped}", heading3_style))
                story.append(Paragraph(f"<b>Engagement:</b> {item.get('engagement', 0):,}", normal_style))
                story.append(Paragraph(f"<b>Engagement Rate:</b> {item.get('engagement_rate', 0):.2f}%", normal_style))
                story.append(Spacer(1, 0.1*inch))
        
        if analytics_report.channel_performance:
            story.append(Paragraph('6.3. Производительность по каналам', heading2_style))
            for channel, perf in analytics_report.channel_performance.items():
                channel_escaped = escape_html(channel)
                story.append(Paragraph(channel_escaped, heading3_style))
                story.append(Paragraph(f"<b>Посты:</b> {perf.get('posts', 0)}", normal_style))
                story.append(Paragraph(f"<b>Просмотры:</b> {perf.get('views', 0):,}", normal_style))
                story.append(Paragraph(f"<b>Вовлеченность:</b> {perf.get('engagement', 0):,}", normal_style))
                story.append(Paragraph(f"<b>Engagement Rate:</b> {perf.get('avg_engagement_rate', 0):.2f}%", normal_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Рекомендации
        if recommendations:
            story.append(Paragraph('6.4. Рекомендации', heading2_style))
            story.append(Spacer(1, 0.15*inch))
            for idx, rec in enumerate(recommendations, 1):
                rec_escaped = escape_html(rec)
                story.append(Paragraph(f"{idx}. {rec_escaped}", bullet_style))
    else:
        story.append(Paragraph("Аналитика не сгенерирована", normal_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Дата создания
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor='#666666',
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}", date_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

# Настройка страницы
st.set_page_config(
    page_title="Content AI Agent",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Минималистичный CSS стиль - поддержка светлой и темной темы
def apply_minimal_style(theme="light"):
    """Применяет минималистичный стиль с поддержкой светлой и темной темы"""
    # Определяем цвета в зависимости от темы
    if theme == "dark":
        bg_color = "#1a1a1a"
        text_color = "#ffffff"
        border_color = "#ffffff"
        input_bg = "#2a2a2a"
        sidebar_bg = "#1a1a1a"
        hover_bg = "rgba(255, 255, 255, 0.1)"
        active_bg = "#333333"
        shadow_color = "rgba(255, 255, 255, 0.1)"
    else:  # light
        bg_color = "#ffffff"
        text_color = "#000000"
        border_color = "#000000"
        input_bg = "#ffffff"
        sidebar_bg = "#ffffff"
        hover_bg = "rgba(0, 0, 0, 0.3)"
        active_bg = "#f0f0f0"
        shadow_color = "rgba(0, 0, 0, 0.1)"
    
    st.markdown(f"""
    <style>
    /* Основной фон */
    .main {{
        padding: 2rem 3rem;
        background-color: {bg_color};
    }}
    
    .stApp {{
        background-color: {bg_color};
    }}
    
    /* Заголовки */
    h1 {{
        font-size: 2.5rem;
        font-weight: 300;
        letter-spacing: -0.02em;
        margin-bottom: 0.5rem;
        color: {text_color};
    }}
    
    h2 {{
        font-size: 1.75rem;
        font-weight: 400;
        letter-spacing: -0.01em;
        margin-top: 2rem;
        margin-bottom: 1rem;
        color: {text_color};
    }}
    
    h3 {{
        font-size: 1.25rem;
        font-weight: 400;
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
        color: {text_color};
    }}
    
    h4 {{
        font-size: 1.1rem;
        font-weight: 400;
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
        color: {text_color};
    }}
    
    /* Текст */
    p, .stMarkdown, div[data-testid="stMarkdownContainer"] {{
        color: {text_color};
        line-height: 1.6;
    }}
    
    /* Все кнопки - прозрачные с обводкой */
    button,
    .stButton > button,
    button[type="submit"],
    button[type="button"],
    .stForm button,
    .stForm .stButton > button {{
        background-color: transparent !important;
        color: {text_color} !important;
        border: 1px solid {border_color} !important;
        border-radius: 2px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 400 !important;
        font-size: 0.95rem !important;
        letter-spacing: 0.01em !important;
        transition: all 0.2s ease !important;
    }}
    
    /* Кнопки с полной шириной */
    .stButton > button,
    .stForm button[type="submit"],
    .stForm .stButton > button {{
        width: 100% !important;
    }}
    
    /* Кнопки при наведении - затемнение по бокам */
    button:hover,
    .stButton > button:hover,
    button[type="submit"]:hover,
    button[type="button"]:hover,
    .stForm button:hover,
    .stForm .stButton > button:hover {{
        background: linear-gradient(to right, 
            {hover_bg} 0%, 
            transparent 20%, 
            transparent 80%, 
            {hover_bg} 100%) !important;
        color: {text_color} !important;
        transform: translateY(-1px);
        border-color: {border_color} !important;
    }}
    
    /* Кнопка "Назад" - меньший размер */
    button[title="Назад"] {{
        padding: 0.5rem 1rem !important;
        width: auto !important;
        font-size: 1rem !important;
    }}
    
    button[title="Назад"]:hover {{
        background: linear-gradient(to right, 
            {hover_bg} 0%, 
            transparent 20%, 
            transparent 80%, 
            {hover_bg} 100%) !important;
        color: {text_color} !important;
        transform: none !important;
    }}
    
    /* Поля ввода */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select {{
        border: 1px solid {border_color};
        border-radius: 2px;
        padding: 0.75rem;
        font-size: 0.95rem;
        background-color: {input_bg};
        color: {text_color} !important;
    }}
    
    /* Убеждаемся, что текст в textarea виден */
    .stTextArea textarea,
    .stTextArea > div > div > textarea,
    textarea[data-testid],
    textarea {{
        color: {text_color} !important;
        background-color: {input_bg} !important;
    }}
    
    /* Для disabled textarea тоже делаем текст видимым */
    .stTextArea textarea:disabled,
    .stTextArea > div > div > textarea:disabled,
    textarea:disabled {{
        color: {text_color} !important;
        background-color: {input_bg} !important;
        -webkit-text-fill-color: {text_color} !important;
        opacity: 1 !important;
    }}
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {{
        border-color: {border_color};
        box-shadow: 0 0 0 2px {shadow_color};
    }}
    
    /* Sidebar */
    [data-testid="stSidebar"] {{
        background-color: {sidebar_bg};
        border-right: 2px solid {border_color};
    }}
    
    [data-testid="stSidebar"] .stButton > button {{
        background-color: transparent !important;
        color: {text_color} !important;
        border: 1px solid {border_color} !important;
        text-align: left;
        padding: 0.75rem 1rem;
        font-weight: 400;
        border-radius: 2px;
    }}
    
    [data-testid="stSidebar"] .stButton > button:hover {{
        background: linear-gradient(to right, 
            {hover_bg} 0%, 
            transparent 20%, 
            transparent 80%, 
            {hover_bg} 100%) !important;
        color: {text_color} !important;
        transform: none;
    }}
    
    [data-testid="stSidebar"] p, 
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] .stMarkdown {{
        color: {text_color};
    }}
    
    /* Универсальный селектор для всех кнопок в области header/sidebar collapse */
    /* ВСЕ кнопки в header - делаем круглыми */
    header button,
    header > button,
    header > div > button,
    header > div button,
    header * > button,
    [role="banner"] button,
    [role="banner"] > button,
    .stApp header button,
    .stApp header > button,
    .stApp [role="banner"] button,
    .stApp [role="banner"] > button {{
        border-radius: 50% !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        padding: 8px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }}
    
    /* Кнопки открытия сайдбара - белый кружочек */
    header button:has(svg),
    header > div > button:has(svg),
    header > div button:has(svg),
    [data-testid*="header"] button:has(svg),
    [data-testid*="Header"] button:has(svg),
    [class*="header"] button:has(svg),
    [class*="Header"] button:has(svg),
    [role="banner"] button:has(svg),
    .stApp > header button:has(svg),
    .stApp > [role="banner"] button:has(svg),
    header > button:has(svg),
    [role="banner"] > button:has(svg),
    .stApp header button:has(svg),
    .stApp [role="banner"] button:has(svg),
    button[aria-label*="menu"]:has(svg),
    button[aria-label*="Menu"]:has(svg),
    button[title*="menu"]:has(svg),
    button[title*="Menu"]:has(svg),
    button[aria-label*="sidebar"]:has(svg),
    button[aria-label*="Sidebar"]:has(svg),
    button[title*="sidebar"]:has(svg),
    button[title*="Sidebar"]:has(svg),
    /* Универсальный селектор для всех кнопок в header с SVG - белый фон */
    header button[class*="button"]:has(svg),
    header > * button:has(svg),
    .stApp header > * button:has(svg) {{
        background-color: #ffffff !important;
        background: #ffffff !important;
        border: 1.5px solid #e0e0e0 !important;
        border-radius: 50% !important;
        outline: none !important;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1) !important;
        padding: 8px !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.2s ease !important;
    }}
    
    /* Дополнительный селектор для гарантированного белого фона кнопки открытия сайдбара */
    header button:has(svg),
    header > * > button:has(svg),
    .stApp header button:has(svg),
    .stApp header > * > button:has(svg) {{
        background: #ffffff !important;
        background-color: #ffffff !important;
        background-image: none !important;
    }}
    
    /* SVG внутри кнопок открытия сайдбара - черные стрелки на белом фоне */
    header button:has(svg) svg,
    header > div > button:has(svg) svg,
    header > div button:has(svg) svg,
    [data-testid*="header"] button:has(svg) svg,
    [class*="header"] button:has(svg) svg,
    .stApp header button:has(svg) svg,
    .stApp [role="banner"] button:has(svg) svg {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        color: {text_color} !important;
    }}
    
    header button:has(svg) path,
    header button:has(svg) g,
    header button:has(svg) * {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
    }}
    
    header button:has(svg):hover,
    header > div > button:has(svg):hover,
    header > div button:has(svg):hover,
    [data-testid*="header"] button:has(svg):hover,
    [data-testid*="Header"] button:has(svg):hover,
    [class*="header"] button:has(svg):hover,
    [class*="Header"] button:has(svg):hover,
    [role="banner"] button:has(svg):hover,
    .stApp > header button:has(svg):hover,
    .stApp > [role="banner"] button:has(svg):hover,
    header > button:has(svg):hover,
    [role="banner"] > button:has(svg):hover,
    .stApp header button:has(svg):hover,
    .stApp [role="banner"] button:has(svg):hover,
    button[aria-label*="menu"]:has(svg):hover,
    button[aria-label*="Menu"]:has(svg):hover,
    button[title*="menu"]:has(svg):hover,
    button[title*="Menu"]:has(svg):hover,
    button[aria-label*="sidebar"]:has(svg):hover,
    button[aria-label*="Sidebar"]:has(svg):hover,
    button[title*="sidebar"]:has(svg):hover,
    button[title*="Sidebar"]:has(svg):hover {{
        background-color: #f5f5f5 !important;
        border: 1.5px solid #d0d0d0 !important;
        border-radius: 50% !important;
        outline: none !important;
        transform: scale(1.05);
    }}
    
    /* Все кнопки в header - делаем круглыми (универсальный селектор) */
    header button,
    header > div > button,
    header > div button,
    [data-testid*="header"] button,
    [data-testid*="Header"] button,
    [class*="header"] button,
    [class*="Header"] button,
    [role="banner"] button,
    .stApp > header button,
    .stApp > [role="banner"] button,
    header > button,
    [role="banner"] > button,
    .stApp header button,
    .stApp [role="banner"] button {{
        border-radius: 50% !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        padding: 8px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }}
    
    header button:focus,
    header > div > button:focus,
    header > div button:focus,
    [data-testid*="header"] button:focus,
    [data-testid*="Header"] button:focus,
    [class*="header"] button:focus,
    [class*="Header"] button:focus,
    [role="banner"] button:focus,
    .stApp > header button:focus,
    .stApp > [role="banner"] button:focus,
    header > button:focus,
    [role="banner"] > button:focus,
    .stApp header button:focus,
    .stApp [role="banner"] button:focus,
    header * button:focus,
    [role="banner"] * button:focus,
    .stApp header * button:focus,
    .stApp [role="banner"] * button:focus {{
        background-color: transparent !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
    }}
    
    /* Кнопка сворачивания сайдбара - круглая обводка */
    [data-testid="stSidebarCollapseButton"],
    button[data-testid="stSidebarCollapseButton"],
    [data-testid="stSidebarCollapseButton"] button,
    .stSidebarCollapseButton,
    button.stSidebarCollapseButton,
    button[aria-label*="Close"],
    button[aria-label*="close"],
    button[title*="Close"],
    button[title*="close"] {{
        background-color: transparent !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        color: {text_color} !important;
        box-shadow: none !important;
        padding: 8px !important;
        margin: 0 !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.2s ease !important;
    }}
    
    [data-testid="stSidebarCollapseButton"]:hover,
    button[data-testid="stSidebarCollapseButton"]:hover,
    [data-testid="stSidebarCollapseButton"]:focus,
    button[data-testid="stSidebarCollapseButton"]:focus,
    button[aria-label*="Close"]:hover,
    button[aria-label*="close"]:hover,
    button[title*="Close"]:hover,
    button[title*="close"]:hover {{
        background-color: {hover_bg} !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        transform: scale(1.05);
    }}
    
    [data-testid="stSidebarCollapseButton"] svg,
    [data-testid="stSidebarCollapseButton"] path,
    [data-testid="stSidebarCollapseButton"] g,
    [data-testid="stSidebarCollapseButton"] *,
    .stSidebarCollapseButton svg,
    .stSidebarCollapseButton path,
    .stSidebarCollapseButton g,
    .stSidebarCollapseButton * {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 2 !important;
        color: {text_color} !important;
    }}
    
    /* Увеличиваем размер SVG для лучшей видимости */
    [data-testid="stSidebarCollapseButton"] svg,
    .stSidebarCollapseButton svg {{
        width: 20px !important;
        height: 20px !important;
        min-width: 20px !important;
        min-height: 20px !important;
    }}
    
    /* Альтернативные селекторы для кнопки сворачивания - круглая обводка */
    button[kind="header"],
    [kind="header"],
    button[class*="collapse"],
    button[class*="Collapse"],
    button[class*="sidebar"],
    button[class*="Sidebar"] {{
        background-color: transparent !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        box-shadow: none !important;
        padding: 8px !important;
        margin: 0 !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.2s ease !important;
    }}
    
    button[kind="header"]:hover,
    [kind="header"]:hover,
    button[kind="header"]:focus,
    [kind="header"]:focus,
    button[class*="collapse"]:hover,
    button[class*="Collapse"]:hover,
    button[class*="sidebar"]:hover,
    button[class*="Sidebar"]:hover {{
        background-color: {hover_bg} !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        transform: scale(1.05);
    }}
    
    button[kind="header"] svg,
    button[kind="header"] path,
    button[kind="header"] g,
    button[kind="header"] *,
    [kind="header"] svg,
    [kind="header"] path,
    [kind="header"] g,
    [kind="header"] * {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 2 !important;
        color: {text_color} !important;
    }}
    
    button[kind="header"] svg,
    [kind="header"] svg {{
        width: 20px !important;
        height: 20px !important;
        min-width: 20px !important;
        min-height: 20px !important;
    }}
    
    /* Все SVG элементы в кнопках */
    button svg,
    button path,
    button g {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 1.5 !important;
        color: {text_color} !important;
    }}
    
    /* Специфичные стили для иконок стрелок */
    svg[viewBox*="0 0"] path,
    svg path[d*="M"],
    svg path[d*="m"] {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 2 !important;
        stroke-linecap: round !important;
        stroke-linejoin: round !important;
    }}
    
    /* Дополнительные стили для всех кнопок с иконками - круглая обводка */
    button[aria-label*="collapse"],
    button[aria-label*="Collapse"],
    button[title*="collapse"],
    button[title*="Collapse"],
    button[aria-label*="expand"],
    button[aria-label*="Expand"],
    button[title*="expand"],
    button[title*="Expand"],
    /* Все кнопки в header с обводкой - делаем круглыми */
    header button[style*="border"],
    header button[class*="button"],
    .stApp header button,
    .stApp [role="banner"] button {{
        background-color: transparent !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        padding: 8px !important;
        margin: 0 !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.2s ease !important;
    }}
    
    button[aria-label*="collapse"]:hover,
    button[aria-label*="Collapse"]:hover,
    button[title*="collapse"]:hover,
    button[title*="Collapse"]:hover,
    button[aria-label*="expand"]:hover,
    button[aria-label*="Expand"]:hover,
    button[title*="expand"]:hover,
    button[title*="Expand"]:hover {{
        background-color: {hover_bg} !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        transform: scale(1.05);
    }}
    
    /* Все кнопки с SVG внутри (иконки) в header - круглая обводка */
    header button:has(svg),
    [role="banner"] button:has(svg),
    .stApp header button:has(svg),
    .stApp [role="banner"] button:has(svg),
    header > div button:has(svg),
    header > div > button:has(svg),
    [data-testid*="header"] button:has(svg),
    [class*="header"] button:has(svg) {{
        background-color: transparent !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        padding: 8px !important;
        margin: 0 !important;
        width: 36px !important;
        height: 36px !important;
        min-width: 36px !important;
        min-height: 36px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.2s ease !important;
    }}
    
    header button:has(svg):hover,
    [role="banner"] button:has(svg):hover,
    .stApp header button:has(svg):hover,
    .stApp [role="banner"] button:has(svg):hover,
    header > div button:has(svg):hover,
    header > div > button:has(svg):hover,
    [data-testid*="header"] button:has(svg):hover,
    [class*="header"] button:has(svg):hover {{
        background-color: {hover_bg} !important;
        border: 1.5px solid {border_color} !important;
        border-radius: 50% !important;
        outline: none !important;
        transform: scale(1.05);
    }}
    
    /* Альтернативный селектор для браузеров без поддержки :has() - применяем к родителю через JS или через более специфичные селекторы */
    header button svg,
    [role="banner"] button svg,
    .stApp header button svg,
    .stApp [role="banner"] button svg {{
        /* SVG стили уже определены выше */
    }}
    
    button[aria-label*="collapse"]:hover,
    button[aria-label*="Collapse"]:hover,
    button[title*="collapse"]:hover,
    button[title*="Collapse"]:hover,
    button[aria-label*="expand"]:hover,
    button[aria-label*="Expand"]:hover,
    button[title*="expand"]:hover,
    button[title*="Expand"]:hover,
    button[aria-label*="collapse"]:focus,
    button[aria-label*="Collapse"]:focus,
    button[title*="collapse"]:focus,
    button[title*="Collapse"]:focus {{
        background-color: transparent !important;
        border: none !important;
        border-width: 0 !important;
        outline: none !important;
    }}
    
    button[aria-label*="collapse"] svg,
    button[aria-label*="Collapse"] svg,
    button[title*="collapse"] svg,
    button[title*="Collapse"] svg {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 2.5 !important;
    }}
    
    /* Метрики */
    [data-testid="stMetricValue"] {{
        font-size: 2rem;
        font-weight: 300;
        color: {text_color};
    }}
    
    [data-testid="stMetricLabel"] {{
        color: {text_color};
    }}
    
    /* Разделители */
    hr {{
        border: none;
        border-top: 1px solid {border_color};
        margin: 2rem 0;
    }}
    
    /* Формы */
    .stForm {{
        border: none;
        padding: 0;
        background-color: {bg_color};
    }}
    
    /* Убираем лишние отступы */
    .element-container {{
        margin-bottom: 1.5rem;
    }}
    
    /* Expander */
    .streamlit-expanderHeader {{
        font-weight: 400;
        color: {text_color};
    }}
    
    .streamlit-expanderContent {{
        color: {text_color};
    }}
    
    /* Уведомления - минималистичные, без вертикальных полосок, с цветным фоном */
    .stAlert,
    [data-testid="stAlert"],
    [data-baseweb="notification"] {{
        border-radius: 2px;
        border-left: none !important;
        border: none !important;
        background-color: {bg_color};
        color: {text_color};
    }}
    
    /* Success - зеленый фон */
    .stSuccess,
    [data-testid="stSuccess"],
    div[data-testid="stSuccess"] {{
        border-left: none !important;
        border: none !important;
        color: {text_color};
    }}
    
    /* Info - синий фон */
    .stInfo,
    [data-testid="stInfo"],
    div[data-testid="stInfo"] {{
        border-left: none !important;
        border: none !important;
        color: {text_color};
    }}
    
    /* Warning - желтый фон */
    .stWarning,
    [data-testid="stWarning"],
    div[data-testid="stWarning"] {{
        border-left: none !important;
        border: none !important;
        color: {text_color};
    }}
    
    /* Error - красный фон */
    .stError,
    [data-testid="stError"],
    div[data-testid="stError"] {{
        border-left: none !important;
        border: none !important;
        color: {text_color};
    }}
    
    /* Цветной фон для success и info сообщений - универсальные селекторы */
    /* Success - зеленый фон */
    div[data-testid="stSuccess"],
    div[data-testid="stSuccess"] > div,
    div[data-testid="stSuccess"] > div > div,
    div[data-testid="stSuccess"] > div > div > div,
    [data-baseweb="notification"][data-testid="stSuccess"],
    [data-baseweb="notification"][data-testid="stSuccess"] > div,
    [data-baseweb="notification"][data-testid="stSuccess"] > div > div,
    .element-container:has([data-testid="stSuccess"]),
    .element-container:has([data-testid="stSuccess"]) > div,
    .element-container:has([data-testid="stSuccess"]) > div > div,
    div:has([data-testid="stSuccess"]),
    div:has([data-testid="stSuccess"]) > div {{
        background-color: #d4edda !important;
    }}
    
    /* Info - синий фон */
    div[data-testid="stInfo"],
    div[data-testid="stInfo"] > div,
    div[data-testid="stInfo"] > div > div,
    div[data-testid="stInfo"] > div > div > div,
    [data-baseweb="notification"][data-testid="stInfo"],
    [data-baseweb="notification"][data-testid="stInfo"] > div,
    [data-baseweb="notification"][data-testid="stInfo"] > div > div,
    .element-container:has([data-testid="stInfo"]),
    .element-container:has([data-testid="stInfo"]) > div,
    .element-container:has([data-testid="stInfo"]) > div > div,
    div:has([data-testid="stInfo"]),
    div:has([data-testid="stInfo"]) > div {{
        background-color: #d1ecf1 !important;
    }}
    
    /* Warning - желтый фон */
    div[data-testid="stWarning"],
    div[data-testid="stWarning"] > div,
    div[data-testid="stWarning"] > div > div,
    div[data-testid="stWarning"] > div > div > div,
    [data-baseweb="notification"][data-testid="stWarning"],
    [data-baseweb="notification"][data-testid="stWarning"] > div,
    [data-baseweb="notification"][data-testid="stWarning"] > div > div,
    .element-container:has([data-testid="stWarning"]),
    .element-container:has([data-testid="stWarning"]) > div,
    .element-container:has([data-testid="stWarning"]) > div > div,
    div:has([data-testid="stWarning"]),
    div:has([data-testid="stWarning"]) > div {{
        background-color: #fff3cd !important;
    }}
    
    /* Error - красный фон */
    div[data-testid="stError"],
    div[data-testid="stError"] > div,
    div[data-testid="stError"] > div > div,
    div[data-testid="stError"] > div > div > div,
    [data-baseweb="notification"][data-testid="stError"],
    [data-baseweb="notification"][data-testid="stError"] > div,
    [data-baseweb="notification"][data-testid="stError"] > div > div,
    .element-container:has([data-testid="stError"]),
    .element-container:has([data-testid="stError"]) > div,
    .element-container:has([data-testid="stError"]) > div > div,
    div:has([data-testid="stError"]),
    div:has([data-testid="stError"]) > div {{
        background-color: #f8d7da !important;
    }}
    
    /* Дополнительные селекторы для полного скрытия полосок */
    div[data-testid="stAlert"] > div,
    div[data-testid="stSuccess"] > div,
    div[data-testid="stInfo"] > div,
    div[data-testid="stWarning"] > div,
    div[data-testid="stError"] > div {{
        border-left: none !important;
        border: none !important;
    }}
    
    /* Убираем псевдоэлементы ::before, которые могут создавать полоски */
    .stAlert::before,
    .stSuccess::before,
    .stInfo::before,
    .stWarning::before,
    .stError::before,
    [data-testid="stAlert"]::before,
    [data-testid="stSuccess"]::before,
    [data-testid="stInfo"]::before,
    [data-testid="stWarning"]::before,
    [data-testid="stError"]::before {{
        display: none !important;
        content: none !important;
    }}
    
    /* Дополнительные селекторы для применения фона через все возможные пути */
    [data-testid="stSuccess"] [data-baseweb="notification"],
    [data-testid="stInfo"] [data-baseweb="notification"],
    [data-testid="stWarning"] [data-baseweb="notification"],
    [data-testid="stError"] [data-baseweb="notification"] {{
        background-color: inherit !important;
    }}
    
    /* Применяем фон через все элементы внутри контейнеров */
    div:has([data-testid="stSuccess"]) [data-baseweb="notification"],
    div:has([data-testid="stInfo"]) [data-baseweb="notification"],
    div:has([data-testid="stWarning"]) [data-baseweb="notification"],
    div:has([data-testid="stError"]) [data-baseweb="notification"] {{
        background-color: inherit !important;
    }}
    </style>
    
    <script>
    // Применяем фон к сообщениям через JavaScript для гарантированного применения
    function applyAlertBackgrounds() {{
        // Success - зеленый фон
        const successElements = document.querySelectorAll('[data-testid="stSuccess"]');
        successElements.forEach(el => {{
            // Применяем фон к самому элементу
            el.style.setProperty('background-color', '#d4edda', 'important');
            // Находим родительский контейнер и применяем фон
            let parent = el.parentElement;
            while (parent && parent !== document.body) {{
                if (parent.classList.contains('element-container') || 
                    parent.hasAttribute('data-testid') || 
                    parent.querySelector('[data-testid="stSuccess"]')) {{
                    parent.style.setProperty('background-color', '#d4edda', 'important');
                }}
                parent = parent.parentElement;
            }}
            // Применяем ко всем вложенным элементам
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                child.style.setProperty('background-color', '#d4edda', 'important');
            }});
        }});
        
        // Info - синий фон
        const infoElements = document.querySelectorAll('[data-testid="stInfo"]');
        infoElements.forEach(el => {{
            // Применяем фон к самому элементу
            el.style.setProperty('background-color', '#d1ecf1', 'important');
            // Находим родительский контейнер и применяем фон
            let parent = el.parentElement;
            while (parent && parent !== document.body) {{
                if (parent.classList.contains('element-container') || 
                    parent.hasAttribute('data-testid') || 
                    parent.querySelector('[data-testid="stInfo"]')) {{
                    parent.style.setProperty('background-color', '#d1ecf1', 'important');
                }}
                parent = parent.parentElement;
            }}
            // Применяем ко всем вложенным элементам
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                child.style.setProperty('background-color', '#d1ecf1', 'important');
            }});
        }});
        
        // Warning - желтый фон
        const warningElements = document.querySelectorAll('[data-testid="stWarning"]');
        warningElements.forEach(el => {{
            el.style.setProperty('background-color', '#fff3cd', 'important');
            let parent = el.parentElement;
            while (parent && parent !== document.body) {{
                if (parent.classList.contains('element-container') || 
                    parent.hasAttribute('data-testid') || 
                    parent.querySelector('[data-testid="stWarning"]')) {{
                    parent.style.setProperty('background-color', '#fff3cd', 'important');
                }}
                parent = parent.parentElement;
            }}
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                child.style.setProperty('background-color', '#fff3cd', 'important');
            }});
        }});
        
        // Error - красный фон
        const errorElements = document.querySelectorAll('[data-testid="stError"]');
        errorElements.forEach(el => {{
            el.style.setProperty('background-color', '#f8d7da', 'important');
            let parent = el.parentElement;
            while (parent && parent !== document.body) {{
                if (parent.classList.contains('element-container') || 
                    parent.hasAttribute('data-testid') || 
                    parent.querySelector('[data-testid="stError"]')) {{
                    parent.style.setProperty('background-color', '#f8d7da', 'important');
                }}
                parent = parent.parentElement;
            }}
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                child.style.setProperty('background-color', '#f8d7da', 'important');
            }});
        }});
    }}
    
    // Применяем при загрузке страницы
    if (document.readyState === 'loading') {{
        document.addEventListener('DOMContentLoaded', applyAlertBackgrounds);
    }} else {{
        applyAlertBackgrounds();
    }}
    
    // Применяем после обновления Streamlit (используем MutationObserver)
    const observer = new MutationObserver(function(mutations) {{
        applyAlertBackgrounds();
    }});
    
    observer.observe(document.body, {{
        childList: true,
        subtree: true
    }});
    
    // Также применяем с задержкой для гарантии
    setTimeout(applyAlertBackgrounds, 100);
    setTimeout(applyAlertBackgrounds, 500);
    setTimeout(applyAlertBackgrounds, 1000);
    </script>
    
    <style>
    /* Caption */
    .stCaption {{
        color: {text_color};
    }}
    
    /* Пространство */
    .block-container {{
        padding-top: 3rem;
        padding-bottom: 3rem;
        background-color: {bg_color};
    }}
    
    /* Табы */
    .stTabs [data-baseweb="tab-list"] {{
        border-bottom: 1px solid {border_color};
    }}
    
    .stTabs [data-baseweb="tab"] {{
        color: {text_color};
    }}
    
    .stTabs [aria-selected="true"] {{
        color: {text_color};
        border-bottom: 2px solid {border_color};
    }}
    
    /* Selectbox и другие элементы */
    .stSelectbox label,
    .stTextInput label,
    .stTextArea label,
    .stNumberInput label {{
        color: {text_color};
    }}
    
    /* Multiselect */
    .stMultiSelect label {{
        color: {text_color};
    }}
    
    /* Slider */
    .stSlider label {{
        color: {text_color};
    }}
    
    /* Убираем дефолтные тени и градиенты */
    .main .block-container {{
        background-color: {bg_color};
    }}
    
    /* Цвет ссылок */
    a {{
        color: {text_color};
        text-decoration: underline;
    }}
    
    a:hover {{
        opacity: 0.8;
    }}
    
    /* Знаки вопроса (help icons) - делаем видимыми на белом фоне */
    .stTooltipIcon,
    [data-testid="stTooltipIcon"],
    button[aria-label*="help"],
    button[aria-label*="Help"],
    .stMarkdown [data-testid="stTooltipIcon"],
    .stTextInput [data-testid="stTooltipIcon"],
    .stTextArea [data-testid="stTooltipIcon"],
    .stSelectbox [data-testid="stTooltipIcon"],
    .stNumberInput [data-testid="stTooltipIcon"],
    .stSlider [data-testid="stTooltipIcon"],
    .stDateInput [data-testid="stTooltipIcon"],
    .stTimeInput [data-testid="stTooltipIcon"],
    .stFileUploader [data-testid="stTooltipIcon"],
    .stColorPicker [data-testid="stTooltipIcon"],
    /* Более универсальные селекторы */
    [class*="stTooltip"],
    [class*="stTooltipIcon"],
    [class*="tooltip-icon"],
    [class*="help-icon"] {{
        color: {text_color} !important;
        fill: {text_color} !important;
        stroke: {text_color} !important;
        opacity: 1 !important;
    }}
    
    .stTooltipIcon:hover,
    [data-testid="stTooltipIcon"]:hover,
    button[aria-label*="help"]:hover,
    button[aria-label*="Help"]:hover,
    [class*="stTooltip"]:hover,
    [class*="stTooltipIcon"]:hover,
    [class*="tooltip-icon"]:hover,
    [class*="help-icon"]:hover {{
        color: {text_color} !important;
        fill: {text_color} !important;
        stroke: {text_color} !important;
        opacity: 1 !important;
    }}
    
    /* SVG внутри help icons - делаем знак вопроса более видимым */
    .stTooltipIcon svg,
    [data-testid="stTooltipIcon"] svg,
    [class*="stTooltip"] svg,
    [class*="stTooltipIcon"] svg,
    [class*="tooltip-icon"] svg,
    [class*="help-icon"] svg {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        color: {text_color} !important;
        width: 18px !important;
        height: 18px !important;
    }}
    
    /* Круг в help icon - делаем более контрастным */
    .stTooltipIcon circle,
    [data-testid="stTooltipIcon"] circle,
    .stTooltipIcon svg circle,
    [data-testid="stTooltipIcon"] svg circle,
    [class*="stTooltip"] circle,
    [class*="stTooltipIcon"] circle,
    .stTooltipIcon svg > circle,
    [data-testid="stTooltipIcon"] svg > circle {{
        fill: transparent !important;
        stroke: {text_color} !important;
        stroke-width: 2 !important;
    }}
    
    /* Знак вопроса внутри круга - делаем более толстым и видимым */
    .stTooltipIcon path,
    [data-testid="stTooltipIcon"] path,
    .stTooltipIcon text,
    [data-testid="stTooltipIcon"] text,
    .stTooltipIcon svg path,
    [data-testid="stTooltipIcon"] svg path,
    .stTooltipIcon svg text,
    [data-testid="stTooltipIcon"] svg text,
    .stTooltipIcon svg > path,
    [data-testid="stTooltipIcon"] svg > path,
    .stTooltipIcon svg > text,
    [data-testid="stTooltipIcon"] svg > text,
    [class*="stTooltip"] path,
    [class*="stTooltipIcon"] path,
    [class*="stTooltip"] text,
    [class*="stTooltipIcon"] text {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        stroke-width: 2 !important;
        font-weight: bold !important;
        font-size: 14px !important;
        font-family: Arial, sans-serif !important;
    }}
    
    /* Если знак вопроса это текст - делаем его жирным */
    .stTooltipIcon svg text,
    [data-testid="stTooltipIcon"] svg text {{
        fill: {text_color} !important;
        stroke: none !important;
        font-weight: 900 !important;
        font-size: 12px !important;
        dominant-baseline: middle !important;
        text-anchor: middle !important;
    }}
    
    /* Группы элементов в SVG */
    .stTooltipIcon g,
    [data-testid="stTooltipIcon"] g,
    [class*="stTooltip"] g,
    [class*="stTooltipIcon"] g {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
    }}
    
    /* Все элементы внутри help icon */
    .stTooltipIcon *,
    [data-testid="stTooltipIcon"] *,
    [class*="stTooltip"] *,
    [class*="stTooltipIcon"] * {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
        color: {text_color} !important;
    }}
    
    /* Альтернативные селекторы для help icons */
    [class*="tooltip"],
    [class*="Tooltip"],
    [class*="help"],
    [class*="Help"] {{
        color: {text_color} !important;
        fill: {text_color} !important;
        stroke: {text_color} !important;
    }}
    
    [class*="tooltip"] svg,
    [class*="Tooltip"] svg,
    [class*="help"] svg,
    [class*="Help"] svg {{
        fill: {text_color} !important;
        stroke: {text_color} !important;
    }}
    </style>
    
    <script>
    // Применяем фон к сообщениям через JavaScript для гарантированного применения
    function applyAlertBackgrounds() {{
        // Success - зеленый фон
        const successElements = document.querySelectorAll('[data-testid="stSuccess"]');
        successElements.forEach(el => {{
            el.style.backgroundColor = '#d4edda';
            // Применяем ко всем вложенным элементам
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                if (child.style) {{
                    child.style.backgroundColor = '#d4edda';
                }}
            }});
        }});
        
        // Info - синий фон
        const infoElements = document.querySelectorAll('[data-testid="stInfo"]');
        infoElements.forEach(el => {{
            el.style.backgroundColor = '#d1ecf1';
            // Применяем ко всем вложенным элементам
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                if (child.style) {{
                    child.style.backgroundColor = '#d1ecf1';
                }}
            }});
        }});
        
        // Warning - желтый фон
        const warningElements = document.querySelectorAll('[data-testid="stWarning"]');
        warningElements.forEach(el => {{
            el.style.backgroundColor = '#fff3cd';
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                if (child.style) {{
                    child.style.backgroundColor = '#fff3cd';
                }}
            }});
        }});
        
        // Error - красный фон
        const errorElements = document.querySelectorAll('[data-testid="stError"]');
        errorElements.forEach(el => {{
            el.style.backgroundColor = '#f8d7da';
            const children = el.querySelectorAll('*');
            children.forEach(child => {{
                if (child.style) {{
                    child.style.backgroundColor = '#f8d7da';
                }}
            }});
        }});
    }}
    
    // Применяем при загрузке страницы
    if (document.readyState === 'loading') {{
        document.addEventListener('DOMContentLoaded', applyAlertBackgrounds);
    }} else {{
        applyAlertBackgrounds();
    }}
    
    // Применяем после обновления Streamlit (используем MutationObserver)
    const observer = new MutationObserver(function(mutations) {{
        applyAlertBackgrounds();
    }});
    
    observer.observe(document.body, {{
        childList: true,
        subtree: true
    }});
    
    // Также применяем с задержкой для гарантии
    setTimeout(applyAlertBackgrounds, 100);
    setTimeout(applyAlertBackgrounds, 500);
    setTimeout(applyAlertBackgrounds, 1000);
    </script>
    """, unsafe_allow_html=True)

# Функция инициализации сессии
def init_session_state():
    """Инициализирует все необходимые переменные сессии"""
    if 'agent' not in st.session_state:
        # Создаем OpenAI клиент, если ключ доступен
        llm_client = create_openai_client()
        st.session_state.agent = ContentAIAgent(llm_client=llm_client)
    if 'current_stage' not in st.session_state:
        st.session_state.current_stage = "onboarding"
    if 'product_info' not in st.session_state:
        st.session_state.product_info = {}
    if 'audience_insights' not in st.session_state:
        st.session_state.audience_insights = {}
    if 'strategy' not in st.session_state:
        st.session_state.strategy = None
    if 'content_list' not in st.session_state:
        st.session_state.content_list = []
    if 'schedule' not in st.session_state:
        st.session_state.schedule = {}
    if 'theme' not in st.session_state:
        st.session_state.theme = "light"
    # Сохранение данных форм для восстановления при возврате
    if 'deep_impact_audit' not in st.session_state:
        st.session_state.deep_impact_audit = {
            "business_niche": "",
            "business_stage": "Идея",
            "existing_data": ""
        }
    if 'strategy_goals' not in st.session_state:
        st.session_state.strategy_goals = "Рост видимости, Привлечение клиентов"
    if 'content_count' not in st.session_state:
        st.session_state.content_count = 10
    if 'scheduling_data' not in st.session_state:
        st.session_state.scheduling_data = {
            "channels": ["linkedin"],
            "start_date": datetime.now().date() + timedelta(days=1),
            "posts_per_week": 5,
            "selected_ids": []
        }
    if 'analytics_days' not in st.session_state:
        st.session_state.analytics_days = 30

# Инициализация сессии при загрузке модуля
init_session_state()

def render_header():
    """Отображение заголовка"""
    init_session_state()
    text_color = "#000000" if st.session_state.theme == "light" else "#ffffff"
    st.markdown("<h1>Content AI Agent</h1>", unsafe_allow_html=True)
    st.markdown(f"<p style='color: {text_color}; font-size: 1rem; margin-bottom: 3rem;'>Автономная система контент-маркетинга для стартапов</p>", unsafe_allow_html=True)
    
    # Блок статуса подключения OpenAI API и кнопка проверки
    status_col, button_col = st.columns([4, 1])
    with status_col:
        # Проверяем переменную окружения напрямую (для актуальности)
        api_key = Config.get_openai_api_key() or os.getenv("OPENAI_API_KEY", "").strip()
        if not api_key:
            st.info("OpenAI API ключ не настроен. Некоторые функции могут быть недоступны.")
            st.info("💡 Установите переменную окружения OPENAI_API_KEY в настройках деплоя.")
        else:
            # Показываем статус по результатам последней проверки
            openai_ok = st.session_state.get("openai_ok", False)
            openai_checked = st.session_state.get("openai_checked", False)
            
            if openai_ok:
                st.success("✅ OpenAI API подключен и работает.")
            elif openai_checked:
                st.warning("⚠️ Ключ OpenAI задан, но последнее подключение завершилось с ошибкой.")
            else:
                st.info("ℹ️ Ключ OpenAI задан. Нажмите кнопку, чтобы проверить подключение.")
    
    with button_col:
        # Проверяем переменную окружения динамически
        if api_key:
            if st.button("Проверить API", key="check_api_status", use_container_width=True):
                # Сбрасываем флаг проверки и принудительно проверяем подключение
                st.session_state.openai_checked = False
                client = create_openai_client()
                
                if client and st.session_state.get("openai_ok"):
                    st.success("✅ OpenAI API подключен и работает.")
                    st.rerun()

def render_sidebar():
    """Боковая панель с навигацией"""
    # Убеждаемся, что сессия инициализирована
    init_session_state()
    
    # Определяем цвета в зависимости от темы
    text_color = "#000000" if st.session_state.theme == "light" else "#ffffff"
    
    # Кнопка переключения темы над "Навигацией"
    theme_icon = "🌙" if st.session_state.theme == "light" else "☀️"
    theme_text = "Темная тема" if st.session_state.theme == "light" else "Светлая тема"
    
    if st.sidebar.button(f"{theme_icon} {theme_text}", key="theme_toggle", use_container_width=True):
        st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"
        st.rerun()
    
    st.sidebar.markdown("<hr style='margin: 1rem 0; border-color: " + text_color + ";'>", unsafe_allow_html=True)
    
    st.sidebar.markdown("<h3 style='font-size: 0.9rem; font-weight: 500; letter-spacing: 0.1em; text-transform: uppercase; color: " + text_color + "; margin-bottom: 2rem;'>Навигация</h3>", unsafe_allow_html=True)
    
    stages = {
        "onboarding": "Онбординг",
        "deep_impact": "DEEP IMPACT",
        "strategy": "Стратегия",
        "content": "Контент",
        "scheduling": "Планирование",
        "analytics": "Аналитика"
    }
    
    current = st.session_state.current_stage
    
    # Определяем цвета для активного/неактивного состояния
    active_bg = "#f0f0f0" if st.session_state.theme == "light" else "#333333"
    
    # Интерактивная навигация
    for stage_key, stage_name in stages.items():
        if stage_key == current:
            st.sidebar.markdown(f"<p style='color: {text_color}; font-weight: 500; padding: 0.75rem 1rem; margin: 0.25rem 0; background-color: {active_bg}; border-left: 2px solid {text_color};'>{stage_name}</p>", unsafe_allow_html=True)
        else:
            # Создаем кнопку для перехода к другому этапу
            if st.sidebar.button(stage_name, key=f"nav_{stage_key}", use_container_width=True):
                st.session_state.current_stage = stage_key
                st.rerun()
    
    st.sidebar.markdown("<hr style='margin: 2rem 0; border-color: " + text_color + ";'>", unsafe_allow_html=True)
    
    # Статус API ключа (динамическая проверка)
    api_key_status = Config.get_openai_api_key() or os.getenv("OPENAI_API_KEY", "").strip()
    api_status = "Подключен" if api_key_status else "Не подключен"
    if st.session_state.theme == "dark":
        api_color = "#ffffff" if api_key_status else "#888888"
    else:
        api_color = "#000000" if api_key_status else "#666666"
    st.sidebar.markdown(f"<p style='font-size: 0.85rem; color: {api_color}; margin: 1rem 0;'>API: {api_status}</p>", unsafe_allow_html=True)
    
    st.sidebar.markdown("<hr style='margin: 1rem 0; border-color: " + text_color + ";'>", unsafe_allow_html=True)
    
    # Текущее состояние
    if st.session_state.agent.product_profile:
        st.sidebar.markdown("<h3 style='font-size: 0.9rem; font-weight: 500; letter-spacing: 0.1em; text-transform: uppercase; color: " + text_color + "; margin-top: 2rem; margin-bottom: 1rem;'>Статус</h3>", unsafe_allow_html=True)
        state = st.session_state.agent.get_current_state()
        st.sidebar.metric("Контент", state['content_count'])
        st.sidebar.metric("Запланировано", state['scheduled_posts'])

def render_onboarding():
    """Экран онбординга"""
    # Кнопка "Назад" - на первой странице не нужна
    st.markdown("<h2>Онбординг</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Привет! Я твой AI-ассистент по контент-маркетингу. Давай начнем с онбординга! "
        "Заполни информацию о своем продукте: название, категорию, описание, функции и целевую аудиторию. "
        "Это поможет мне лучше понять твой продукт и создать эффективную контент-стратегию.",
        theme=st.session_state.theme
    )
    
    st.markdown("<p style='color: #000000; margin-bottom: 2rem;'>Начнем с анализа вашего продукта</p>", unsafe_allow_html=True)
    
    with st.form("product_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            product_name = st.text_input(
                "Название продукта *",
                value=st.session_state.product_info.get("name", ""),
                help="Название вашего продукта или сервиса"
            )
            
            category_options = ["SaaS", "E-commerce", "Образование", "Здоровье", "Финансы", "Другое"]
            saved_category = st.session_state.product_info.get("category", "")
            category_index = category_options.index(saved_category) if saved_category in category_options else 0
            product_category = st.selectbox(
                "Категория *",
                category_options,
                index=category_index
            )
        
        with col2:
            product_description = st.text_area(
                "Описание продукта *",
                value=st.session_state.product_info.get("description", ""),
                height=100,
                help="Краткое описание того, что делает ваш продукт"
            )
        
        features_input = st.text_input(
            "Ключевые функции (через запятую) *",
            value=", ".join(st.session_state.product_info.get("features", [])),
            help="Основные функции и возможности продукта"
        )
        
        target_audience = st.text_input(
            "Целевая аудитория (предварительно)",
            value=st.session_state.product_info.get("target_audience", ""),
            help="Кто ваши потенциальные клиенты?"
        )
        
        submitted = st.form_submit_button("Начать анализ", use_container_width=True)
        
        if submitted:
            if not all([product_name, product_description, features_input]):
                show_error_message("Пожалуйста, заполните все обязательные поля")
            else:
                with st.spinner("Анализируем продукт..."):
                    product_info = {
                        "name": product_name,
                        "description": product_description,
                        "category": product_category,
                        "features": [f.strip() for f in features_input.split(",") if f.strip()],
                        "target_audience": target_audience
                    }
                    
                    st.session_state.product_info = product_info
                    result = st.session_state.agent.start_onboarding(product_info)
                    
                    st.session_state.current_stage = "deep_impact"
                    st.success("Продукт проанализирован")
                    st.rerun()

def render_deep_impact():
    """Экран DEEP IMPACT анализа"""
    # Кнопка "Назад"
    col_back, col_title = st.columns([1, 10])
    with col_back:
        if st.button("←", key="back_deep_impact", help="Назад"):
            st.session_state.current_stage = "onboarding"
            st.rerun()
    with col_title:
        st.markdown("<h2>DEEP IMPACT</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Отлично! Теперь переходим к анализу DEEP IMPACT. "
        "Сначала проведу аудит твоих данных: укажи бизнес-нишу, стадию развития и существующие данные о продукте. "
        "Затем мы сгенерируем гипотезы по уровням: Рынок → Ниша → Аудитория → Сегмент → Боль → Задача. "
        "Это поможет точно определить целевую аудиторию!",
        theme=st.session_state.theme
    )
    
    st.markdown("<p style='color: #000000; margin-bottom: 2rem;'>Проведем глубокий анализ вашей целевой аудитории</p>", unsafe_allow_html=True)
    
    tab1, tab2, tab3 = st.tabs(["Аудит данных", "Генерация гипотез", "Расширение инсайтов"])
    
    with tab1:
        st.markdown("<h3>Аудит исходных данных</h3>", unsafe_allow_html=True)
        
        with st.form("audit_form"):
            business_niche = st.text_input(
                "Бизнес-ниша",
                value=st.session_state.deep_impact_audit.get("business_niche", "")
            )
            stage_options = ["Идея", "MVP", "Первые клиенты", "Рост", "Масштабирование"]
            saved_stage = st.session_state.deep_impact_audit.get("business_stage", "Идея")
            stage_index = stage_options.index(saved_stage) if saved_stage in stage_options else 0
            business_stage = st.selectbox(
                "Стадия развития",
                stage_options,
                index=stage_index
            )
            existing_data = st.text_area(
                "Какие данные о аудитории у вас есть?",
                value=st.session_state.deep_impact_audit.get("existing_data", ""),
                height=100
            )
            
            if st.form_submit_button("Провести аудит", use_container_width=True):
                with st.spinner("Проводим аудит..."):
                    business_context = {
                        "business_niche": business_niche,
                        "stage": business_stage,
                        "existing_data": existing_data
                    }
                    
                    # Сохраняем данные формы
                    st.session_state.deep_impact_audit = {
                        "business_niche": business_niche,
                        "business_stage": business_stage,
                        "existing_data": existing_data
                    }
                    
                    result = st.session_state.agent.run_deep_impact_audit(business_context)
                    
                    # Сохраняем результат аудита
                    st.session_state.audit_result = result
                    
                    # Показываем объединенное сообщение с зеленым фоном
                    show_success_message("Аудит данных завершен. Теперь вы можете перейти к генерации гипотез во вкладке 'Генерация гипотез'.")
    
    with tab2:
        st.markdown("<h3>Генерация гипотез</h3>", unsafe_allow_html=True)
        st.markdown("<p style='color: #000000; margin-bottom: 1.5rem;'>Интерактивный процесс генерации гипотез по уровням: Рынок → Ниша → Аудитория → Сегмент → Боль → Задача</p>", unsafe_allow_html=True)
        
        # Инициализация состояния для гипотез
        if 'hypothesis_levels' not in st.session_state:
            st.session_state.hypothesis_levels = {
                "market": {"completed": False, "selected": None, "hypotheses": []},
                "niche": {"completed": False, "selected": None, "hypotheses": []},
                "audience": {"completed": False, "selected": None, "hypotheses": []},
                "segment": {"completed": False, "selected": None, "hypotheses": []},
                "pain": {"completed": False, "selected": None, "hypotheses": []},
                "task": {"completed": False, "selected": None, "hypotheses": []}
            }
        
        # Названия уровней
        level_names = {
            "market": "Рынок",
            "niche": "Ниша",
            "audience": "Аудитория",
            "segment": "Сегмент",
            "pain": "Боль",
            "task": "Задача"
        }
        
        levels_order = ["market", "niche", "audience", "segment", "pain", "task"]
        
        # Инициализация выбранного уровня для работы
        if 'selected_hypothesis_level' not in st.session_state:
            # Определяем первый незавершенный уровень
            for level in levels_order:
                if not st.session_state.hypothesis_levels[level]["completed"]:
                    st.session_state.selected_hypothesis_level = level
                    break
            else:
                st.session_state.selected_hypothesis_level = levels_order[0]
        
        # Проверяем, все ли уровни завершены - показываем блок скачивания ПЕРЕД прогрессом
        all_completed = all(
            st.session_state.hypothesis_levels[level]["completed"] 
            for level in levels_order
        )
        
        # Если все уровни завершены, показываем блок скачивания
        if all_completed:
            st.markdown("<hr style='margin: 2rem 0; border-color: #000000;'>", unsafe_allow_html=True)
            
            # Центрируем блок
            col_left, col_center, col_right = st.columns([1, 2, 1])
            
            with col_center:
                st.markdown("<h4 style='text-align: center; margin-bottom: 1rem;'>Все уровни завершены</h4>", unsafe_allow_html=True)
                st.markdown("<p style='color: #000000; text-align: center; margin-bottom: 1.5rem;'>Вы можете скачать отчет с выбранными гипотезами</p>", unsafe_allow_html=True)
                
                # Выбор формата
                file_format = st.selectbox(
                    "Формат файла",
                    ["DOCX", "PDF"],
                    key="report_format"
                )
                
                # Кнопка скачивания под форматом
                try:
                    if file_format == "DOCX":
                        if DOCX_AVAILABLE:
                            report_data = generate_docx_report(
                                st.session_state.hypothesis_levels,
                                st.session_state.product_info
                            )
                            st.download_button(
                                label="Скачать отчет (DOCX)",
                                data=report_data,
                                file_name=f"hypotheses_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            show_error_message("Библиотека python-docx не установлена. Установите: pip install python-docx")
                    else:  # PDF
                        if PDF_AVAILABLE:
                            report_data = generate_pdf_report(
                                st.session_state.hypothesis_levels,
                                st.session_state.product_info
                            )
                            st.download_button(
                                label="Скачать отчет (PDF)",
                                data=report_data,
                                file_name=f"hypotheses_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        else:
                            show_error_message("Библиотека reportlab не установлена. Установите: pip install reportlab")
                except Exception as e:
                    show_error_message(f"Ошибка при генерации отчета: {str(e)}")
            
            st.markdown("<hr style='margin: 2rem 0; border-color: #000000;'>", unsafe_allow_html=True)
        
        # Отображаем прогресс с возможностью выбора уровня
        st.markdown("<h4 style='margin-top: 2rem; margin-bottom: 1rem;'>Прогресс</h4>", unsafe_allow_html=True)
        progress_cols = st.columns(len(levels_order))
        for idx, level in enumerate(levels_order):
            with progress_cols[idx]:
                is_completed = st.session_state.hypothesis_levels[level]["completed"]
                is_selected = st.session_state.selected_hypothesis_level == level
                
                if is_selected:
                    status_color = "#000000"
                    status_text = "→"
                    border_style = "border: 2px solid #000000; padding: 0.5rem;"
                elif is_completed:
                    status_color = "#000000"
                    status_text = "✓"
                    border_style = "border: 1px solid #e0e0e0; padding: 0.5rem;"
                else:
                    status_color = "#666666"
                    status_text = "○"
                    border_style = "border: 1px solid #000000; padding: 0.5rem;"
                
                # Кнопка для выбора уровня
                if st.button(level_names[level], key=f"level_btn_{level}", use_container_width=True):
                    st.session_state.selected_hypothesis_level = level
                    st.rerun()
                
                st.markdown(f"<p style='text-align: center; font-size: 0.9rem; color: {status_color}; margin-top: 0.5rem;'>{status_text}</p>", unsafe_allow_html=True)
        
        st.markdown("<hr style='margin: 2rem 0; border-color: #000000;'>", unsafe_allow_html=True)
        
        # Работаем с выбранным уровнем
        current_level = st.session_state.selected_hypothesis_level
        
        if current_level:
            # Отображаем текущий уровень
            st.markdown(f"<h4>Уровень: {level_names[current_level]}</h4>", unsafe_allow_html=True)
            
            # Если уровень завершен, показываем выбранную гипотезу и кнопку "Изменить выбор"
            if st.session_state.hypothesis_levels[current_level]["completed"]:
                selected = st.session_state.hypothesis_levels[current_level]["selected"]
                if selected:
                    st.markdown("<p style='color: #000000; margin-bottom: 1rem;'>Выбранная гипотеза:</p>", unsafe_allow_html=True)
                    st.markdown(f"**{selected.get('name', 'Гипотеза')}**")
                    if selected.get('description'):
                        st.markdown(f"<p style='color: #000000; margin: 0.5rem 0;'>{selected.get('description')}</p>", unsafe_allow_html=True)
                    
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if st.button("Изменить выбор", key=f"change_{current_level}", use_container_width=True):
                            st.session_state.hypothesis_levels[current_level]["completed"] = False
                            st.session_state.hypothesis_levels[current_level]["selected"] = None
                            st.rerun()
                    with col2:
                        if st.button("Сгенерировать заново", key=f"regenerate_{current_level}", use_container_width=True):
                            st.session_state.hypothesis_levels[current_level]["hypotheses"] = []
                            st.session_state.hypothesis_levels[current_level]["completed"] = False
                            st.session_state.hypothesis_levels[current_level]["selected"] = None
                            st.rerun()
                else:
                    st.info("Уровень завершен, но гипотеза не выбрана.")
                    if st.button("Выбрать гипотезу", key=f"select_{current_level}", use_container_width=True):
                        st.session_state.hypothesis_levels[current_level]["completed"] = False
                        st.rerun()
            
            # Если гипотезы еще не сгенерированы
            elif not st.session_state.hypothesis_levels[current_level]["hypotheses"]:
                if st.button(f"Сгенерировать гипотезы для {level_names[current_level]}", use_container_width=True, key=f"gen_{current_level}"):
                    # Показываем прогресс-бар для лучшего UX
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    try:
                        status_text.text("Подготовка контекста...")
                        progress_bar.progress(10)
                        
                        # Получаем контекст
                        context = {}
                        if st.session_state.product_info:
                            context.update(st.session_state.product_info)
                        
                        status_text.text("Анализ предыдущих выборов...")
                        progress_bar.progress(30)
                        
                        # Добавляем предыдущие выборы
                        previous_choices = {}
                        for prev_level in levels_order[:levels_order.index(current_level)]:
                            if st.session_state.hypothesis_levels[prev_level]["selected"]:
                                previous_choices[prev_level] = st.session_state.hypothesis_levels[prev_level]["selected"]
                        context["previous_choices"] = previous_choices
                        
                        status_text.text(f"Генерируем гипотезы для {level_names[current_level]}... (это займет 10-20 секунд)")
                        progress_bar.progress(50)
                        
                        # Генерируем гипотезы
                        llm_client = st.session_state.agent.llm_client if hasattr(st.session_state.agent, 'llm_client') else None
                        result = st.session_state.agent.generate_hypotheses(current_level, context, llm_client=llm_client)
                        
                        status_text.text("Обработка результатов...")
                        progress_bar.progress(90)
                        
                        # Сохраняем гипотезы
                        st.session_state.hypothesis_levels[current_level]["hypotheses"] = result["results"]["hypotheses"]
                        
                        progress_bar.progress(100)
                        status_text.text("Готово!")
                        
                        # Небольшая задержка для показа завершения
                        import time
                        time.sleep(0.5)
                        
                        st.rerun()
                    except Exception as e:
                        status_text.text(f"Ошибка: {str(e)}")
                        progress_bar.progress(0)
                        st.error(f"Произошла ошибка при генерации гипотез: {str(e)}")
            else:
                # Отображаем гипотезы для выбора
                hypotheses = st.session_state.hypothesis_levels[current_level]["hypotheses"]
                
                # Проверяем, не демо-ли это гипотезы (шаблон без LLM)
                is_demo = False
                if hypotheses:
                    first = hypotheses[0]
                    name_sample = str(first.get("name", ""))
                    desc_sample = str(first.get("description", ""))
                    if name_sample.startswith("Гипотеза 1 для") and desc_sample.startswith("Описание гипотезы"):
                        is_demo = True
                
                if is_demo:
                    st.warning(
                        "Похоже, сгенерировались демонстрационные гипотезы по шаблону. "
                        "Проверьте настройку API ключа LLM или попробуйте сгенерировать ещё раз."
                    )
                
                st.markdown(
                    f"<p style='color: #000000; margin-bottom: 1rem;'>Выберите наиболее подходящую гипотезу "
                    f"из {len(hypotheses)} вариантов или перегенерируйте список:</p>",
                    unsafe_allow_html=True
                )
                
                # Кнопка для повторной генерации
                if st.button("Сгенерировать еще раз", key=f"regen_{current_level}", use_container_width=True):
                    # Перегенерация использует тот же контекст, что и первичная
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    try:
                        status_text.text("Подготовка контекста...")
                        progress_bar.progress(10)
                        
                        context = {}
                        if st.session_state.product_info:
                            context.update(st.session_state.product_info)
                        
                        status_text.text("Анализ предыдущих выборов...")
                        progress_bar.progress(30)
                        
                        previous_choices = {}
                        for prev_level in levels_order[:levels_order.index(current_level)]:
                            if st.session_state.hypothesis_levels[prev_level]["selected"]:
                                previous_choices[prev_level] = st.session_state.hypothesis_levels[prev_level]["selected"]
                        context["previous_choices"] = previous_choices
                        
                        status_text.text(f"Генерируем гипотезы для {level_names[current_level]}... (это займет 10-20 секунд)")
                        progress_bar.progress(50)
                        
                        llm_client = st.session_state.agent.llm_client if hasattr(st.session_state.agent, 'llm_client') else None
                        result = st.session_state.agent.generate_hypotheses(current_level, context, llm_client=llm_client)
                        
                        status_text.text("Обработка результатов...")
                        progress_bar.progress(90)
                        
                        st.session_state.hypothesis_levels[current_level]["hypotheses"] = result["results"]["hypotheses"]
                        st.session_state.hypothesis_levels[current_level]["selected"] = None
                        st.session_state.hypothesis_levels[current_level]["completed"] = False
                        
                        progress_bar.progress(100)
                        status_text.text("Готово!")
                        
                        import time
                        time.sleep(0.5)
                        st.rerun()
                    except Exception as e:
                        status_text.text(f"Ошибка: {str(e)}")
                        progress_bar.progress(0)
                        st.error(f"Произошла ошибка при генерации гипотез: {str(e)}")
                
                # Отображаем гипотезы
                selected_index = None
                for idx, hypothesis in enumerate(hypotheses):
                    with st.container():
                        col1, col2 = st.columns([10, 1])
                        with col1:
                            name = hypothesis.get("name", f"Гипотеза {idx+1}")
                            description = hypothesis.get("description", "")
                            characteristics = hypothesis.get("characteristics", [])
                            potential = hypothesis.get("potential", "")
                            
                            st.markdown(f"**{name}**")
                            if description:
                                st.markdown(f"<p style='color: #000000; margin: 0.5rem 0;'>{description}</p>", unsafe_allow_html=True)
                            if characteristics:
                                st.markdown("**Характеристики:**")
                                for char in characteristics[:3]:
                                    st.markdown(f"- {char}")
                            if potential:
                                st.markdown(f"<p style='color: #000000; font-size: 0.9rem; margin-top: 0.5rem;'><em>Потенциал: {potential}</em></p>", unsafe_allow_html=True)
                        with col2:
                            if st.button("Выбрать", key=f"select_{current_level}_{idx}", use_container_width=True):
                                selected_index = idx
                        
                        if idx < len(hypotheses) - 1:
                            st.markdown("<hr style='margin: 1rem 0; border-color: #000000;'>", unsafe_allow_html=True)
                
                # Обработка выбора
                if selected_index is not None:
                    selected_hypothesis = hypotheses[selected_index]
                    st.session_state.hypothesis_levels[current_level]["selected"] = selected_hypothesis
                    st.session_state.hypothesis_levels[current_level]["completed"] = True
                    st.success(f"Выбрана гипотеза: {selected_hypothesis.get('name', 'Гипотеза')}")
                    st.rerun()
                
                # Кнопка для пропуска (если нужно)
                col_skip, col_empty = st.columns([1, 2])
                with col_skip:
                    if st.button("Пропустить этот уровень", key=f"skip_{current_level}", use_container_width=True):
                        st.session_state.hypothesis_levels[current_level]["completed"] = True
                        st.rerun()
    
    with tab3:
        st.markdown("<h3>Расширение инсайтов о аудитории</h3>", unsafe_allow_html=True)
        
        if st.button("Расширить инсайты", use_container_width=True):
            with st.spinner("Расширяем понимание аудитории..."):
                result = st.session_state.agent.expand_audience_insights()
                st.session_state.audience_insights = result['insights']
                st.success("Инсайты расширены")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Главная боль", result['insights'].get('main_pain', 'N/A'))
                    st.metric("Главная задача", result['insights'].get('main_task', 'N/A'))
                
                with col2:
                    if result['insights'].get('motivations'):
                        st.write("**Мотивации:**")
                        for mot in list(result['insights']['motivations'].values())[:3]:
                            st.write(f"- {mot}")
                
                st.session_state.current_stage = "strategy"
                st.rerun()

def render_strategy():
    """Экран генерации стратегии"""
    # Кнопка "Назад"
    col_back, col_title = st.columns([1, 10])
    with col_back:
        if st.button("←", key="back_strategy", help="Назад"):
            st.session_state.current_stage = "deep_impact"
            st.rerun()
    with col_title:
        st.markdown("<h2>Стратегия</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Отлично! Теперь создадим контент-стратегию на основе всех собранных данных. "
        "Укажи бизнес-цели, которые ты хочешь достичь с помощью контента. "
        "Я создам стратегию с контентными столпами, каналами публикации, tone of voice и ключевыми сообщениями.",
        theme=st.session_state.theme
    )
    
    # Проверяем наличие product_profile
    if not st.session_state.agent.product_profile:
        st.warning("Сначала завершите онбординг и заполните информацию о продукте")
        if st.button("Вернуться к онбордингу"):
            st.session_state.current_stage = "onboarding"
            st.rerun()
        return
    
    # Проверяем наличие audience_insights
    if not st.session_state.audience_insights:
        st.warning("Сначала завершите анализ DEEP IMPACT")
        if st.button("Вернуться к DEEP IMPACT"):
            st.session_state.current_stage = "deep_impact"
            st.rerun()
        return
    
    with st.form("strategy_form"):
        business_goals = st.text_input(
            "Бизнес-цели (через запятую)",
            value=st.session_state.strategy_goals,
            help="Какие цели вы хотите достичь с помощью контента?"
        )
        
        if st.form_submit_button("Создать стратегию", use_container_width=True):
            with st.spinner("Генерируем контент-стратегию..."):
                try:
                    goals_list = [g.strip() for g in business_goals.split(",") if g.strip()]
                    
                    # Сохраняем данные формы
                    st.session_state.strategy_goals = business_goals
                    
                    strategy = st.session_state.agent.generate_content_strategy(goals_list)
                    st.session_state.strategy = strategy
                except ValueError as e:
                    show_error_message(f"Ошибка: {str(e)}")
                    if "Product profile not set" in str(e):
                        st.info("Пожалуйста, вернитесь к онбордингу и заполните информацию о продукте")
                        if st.button("Вернуться к онбордингу"):
                            st.session_state.current_stage = "onboarding"
                            st.rerun()
                    elif "Audience insights not set" in str(e):
                        st.info("Пожалуйста, завершите анализ DEEP IMPACT")
                        if st.button("Вернуться к DEEP IMPACT"):
                            st.session_state.current_stage = "deep_impact"
                            st.rerun()
                    return
                
                # Если стратегия успешно создана, показываем результаты
                st.success("Контент-стратегия создана")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("<h3>Контентные столпы</h3>", unsafe_allow_html=True)
                    for i, pillar in enumerate(strategy.content_pillars, 1):
                        st.write(f"{i}. {pillar}")
                
                with col2:
                    st.markdown("<h3>Каналы</h3>", unsafe_allow_html=True)
                    for channel in strategy.channels:
                        st.write(f"- {channel.title()}")
                    
                    st.markdown("<h3>Tone of Voice</h3>", unsafe_allow_html=True)
                    st.write(strategy.tone_of_voice)
                
                st.session_state.current_stage = "content"
                st.rerun()

def render_content():
    """Экран генерации контента"""
    # Кнопка "Назад"
    col_back, col_title = st.columns([1, 10])
    with col_back:
        if st.button("←", key="back_content", help="Назад"):
            st.session_state.current_stage = "strategy"
            st.rerun()
    with col_title:
        st.markdown("<h2>Контент</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Превосходно! Стратегия готова. Теперь сгенерируем контент! "
        "Укажи, сколько постов ты хочешь создать. Я создам контент на основе стратегии, "
        "адаптированный под разные каналы и форматы. Каждый пост будет соответствовать твоим контентным столпам и tone of voice.",
        theme=st.session_state.theme
    )
    
    if not st.session_state.strategy:
        st.warning("Сначала создайте контент-стратегию")
        if st.button("Вернуться к стратегии"):
            st.session_state.current_stage = "strategy"
            st.rerun()
        return
    
    col1, col2 = st.columns([2, 1])
    
    with col2:
        count = st.number_input(
            "Количество постов", 
            min_value=1, 
            max_value=50, 
            value=st.session_state.content_count
        )
        
        if st.button("Сгенерировать контент", use_container_width=True):
            with st.spinner(f"Генерируем {count} постов..."):
                # Сохраняем данные формы
                st.session_state.content_count = count
                
                content_list = st.session_state.agent.generate_content_batch(count)
                st.session_state.content_list = content_list
                st.success(f"Создано {len(content_list)} постов")
                st.rerun()
    
    with col1:
        if st.session_state.content_list:
            st.markdown(f"<h3>Сгенерированный контент ({len(st.session_state.content_list)} постов)</h3>", unsafe_allow_html=True)
            
            for i, content in enumerate(st.session_state.content_list):
                # Используем индекс для гарантированно уникального ключа
                # Индекс начинается с 0, поэтому используем i+1 для отображения, но i для ключей
                content_id = getattr(content, 'id', None)
                unique_suffix = f"{i}_{content_id}" if content_id else str(i)
                
                # Убираем key из expander, так как он не поддерживается в некоторых версиях Streamlit
                with st.expander(f"{content.title} · {content.channel}"):
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.caption(f"Столп: {content.pillar}")
                        st.caption(f"Стадия воронки: {content.funnel_stage}")
                    with col_b:
                        st.caption(f"Канал: {content.channel}")
                        st.caption(f"Статус: {content.status.value}")
                    
                    # Проверяем, что контент не пустой
                    content_text = content.content if content.content else "Контент не был сгенерирован. Попробуйте сгенерировать заново."
                    
                    # Используем markdown для отображения контента с правильным цветом
                    st.markdown(
                        f"""
                        <div style="background-color: {st.session_state.theme == 'dark' and '#2a2a2a' or '#ffffff'}; 
                                    color: {st.session_state.theme == 'dark' and '#ffffff' or '#000000'}; 
                                    padding: 1rem; 
                                    border: 1px solid {st.session_state.theme == 'dark' and '#ffffff' or '#000000'}; 
                                    border-radius: 4px; 
                                    min-height: 200px; 
                                    white-space: pre-wrap; 
                                    font-family: monospace;
                                    font-size: 0.95rem;
                                    line-height: 1.6;">
                        {content_text.replace('<', '&lt;').replace('>', '&gt;')}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
            
            if st.button("Перейти к планированию", use_container_width=True):
                st.session_state.current_stage = "scheduling"
                st.rerun()

def render_scheduling():
    """Экран планирования"""
    # Кнопка "Назад"
    col_back, col_title = st.columns([1, 10])
    with col_back:
        if st.button("←", key="back_scheduling", help="Назад"):
            st.session_state.current_stage = "content"
            st.rerun()
    with col_title:
        st.markdown("<h2>Планирование</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Контент готов! Теперь создадим расписание публикаций. "
        "Выбери каналы, на которых хочешь публиковать, дату начала, количество постов в неделю и конкретные посты. "
        "Я создам оптимальный календарь публикаций для максимального охвата аудитории.",
        theme=st.session_state.theme
    )
    
    if not st.session_state.content_list:
        st.warning("Сначала сгенерируйте контент")
        if st.button("Вернуться к генерации"):
            st.session_state.current_stage = "content"
            st.rerun()
        return
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("<h3>Настройки расписания</h3>", unsafe_allow_html=True)
        
        saved_channels = st.session_state.scheduling_data.get("channels", ["linkedin"])
        channels = st.multiselect(
            "Каналы публикации",
            ["linkedin", "twitter", "telegram"],
            default=saved_channels
        )
        
        saved_start_date = st.session_state.scheduling_data.get("start_date", datetime.now().date() + timedelta(days=1))
        start_date = st.date_input(
            "Дата начала публикаций",
            value=saved_start_date
        )
        
        saved_posts_per_week = st.session_state.scheduling_data.get("posts_per_week", 5)
        posts_per_week = st.slider("Постов в неделю", 1, 10, saved_posts_per_week)
    
    with col2:
        st.markdown("<h3>Выбор контента</h3>", unsafe_allow_html=True)
        st.info(f"Доступно постов: {len(st.session_state.content_list)}")
        
        saved_selected_ids = st.session_state.scheduling_data.get("selected_ids", [])
        # Если сохраненные ID не существуют в текущем списке контента, используем все доступные
        available_ids = [c.id for c in st.session_state.content_list]
        default_ids = [id for id in saved_selected_ids if id in available_ids] if saved_selected_ids else available_ids
        
        selected_ids = st.multiselect(
            "Выберите контент для публикации",
            options=available_ids,
            default=default_ids,
            format_func=lambda x: next(
                (c.title for c in st.session_state.content_list if c.id == x), x
            )
        )
    
    if st.button("Создать расписание", use_container_width=True):
        if not channels:
            show_error_message("Выберите хотя бы один канал")
        elif not selected_ids:
            show_error_message("Выберите контент для публикации")
        else:
            with st.spinner("Создаем расписание..."):
                # Сохраняем данные формы
                st.session_state.scheduling_data = {
                    "channels": channels,
                    "start_date": start_date,
                    "posts_per_week": posts_per_week,
                    "selected_ids": selected_ids
                }
                
                schedule = st.session_state.agent.approve_and_schedule_content(
                    content_ids=selected_ids,
                    start_date=datetime.combine(start_date, datetime.min.time()),
                    channels=channels
                )
                
                st.session_state.schedule = schedule
                st.success("Расписание создано")
                
                st.markdown("<h3>Календарь публикаций</h3>", unsafe_allow_html=True)
                
                for channel, posts in schedule.items():
                    with st.expander(f"{channel.upper()} - {len(posts)} постов"):
                        for post in posts[:10]:  # Показываем первые 10
                            st.write(f"**{post.scheduled_time.strftime('%Y-%m-%d %H:%M')}** - {post.content.title}")
                        if len(posts) > 10:
                            st.caption(f"... и еще {len(posts) - 10} постов")
                
                st.session_state.current_stage = "analytics"
                st.rerun()

def render_analytics():
    """Экран аналитики"""
    # Кнопка "Назад"
    col_back, col_title = st.columns([1, 10])
    with col_back:
        if st.button("←", key="back_analytics", help="Назад"):
            st.session_state.current_stage = "scheduling"
            st.rerun()
    with col_title:
        st.markdown("<h2>Аналитика</h2>", unsafe_allow_html=True)
    
    # Показываем AI-ассистента
    show_ai_assistant(
        "Отлично! Все готово к публикации. Теперь давай посмотрим на аналитику! "
        "Выбери период анализа, и я покажу метрики: количество постов, просмотры, вовлеченность, engagement rate. "
        "Также увидишь топ контент, производительность по каналам и получишь рекомендации по улучшению.",
        theme=st.session_state.theme
    )
    
    days = st.slider("Период анализа (дней)", 7, 90, st.session_state.analytics_days)
    
    # Сохраняем данные формы при изменении слайдера
    if days != st.session_state.analytics_days:
        st.session_state.analytics_days = days
    
    # Показываем отчет, если контент есть, или при нажатии кнопки
    show_report = False
    if st.button("Сгенерировать отчет", use_container_width=True):
        show_report = True
        st.session_state.analytics_days = days
    
    # Также показываем отчет автоматически, если есть сохраненный отчет
    if 'analytics_report' in st.session_state and st.session_state.analytics_report:
        show_report = True
    
    if show_report:
        try:
            with st.spinner("Анализируем данные..."):
                # Передаем список контента из session_state в агент для анализа
                content_list_for_analytics = st.session_state.content_list if st.session_state.content_list else []
                
                # Добавляем контент в историю агента для анализа
                if content_list_for_analytics:
                    for content in content_list_for_analytics:
                        if content not in st.session_state.agent.content_agent.content_history:
                            st.session_state.agent.content_agent.content_history.append(content)
                
                report = st.session_state.agent.get_analytics_report(days=days, content_list=content_list_for_analytics)
                
                # Сохраняем отчет в session_state
                st.session_state.analytics_report = report
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("Всего постов", report.total_posts)
                with col2:
                    st.metric("Просмотры", report.total_views)
                with col3:
                    st.metric("Вовлеченность", report.total_engagement)
                with col4:
                    st.metric("Engagement Rate", f"{report.avg_engagement_rate:.2f}%")
                
                if report.top_performing_content:
                    st.markdown("<h3>Топ контент</h3>", unsafe_allow_html=True)
                    for item in report.top_performing_content:
                        col_a, col_b = st.columns([3, 1])
                        with col_a:
                            st.write(f"**{item.get('title', 'Без названия')}**")
                        with col_b:
                            st.metric("Engagement", item.get('engagement', 0))
                else:
                    st.info("Пока нет контента для анализа. Сгенерируйте и опубликуйте контент, чтобы увидеть аналитику.")
                
                if report.channel_performance:
                    st.markdown("<h3>Производительность по каналам</h3>", unsafe_allow_html=True)
                    channel_data = []
                    for channel, perf in report.channel_performance.items():
                        channel_data.append({
                            "Канал": channel,
                            "Посты": perf.get("posts", 0),
                            "Просмотры": perf.get("views", 0),
                            "Engagement Rate": f"{perf.get('avg_engagement_rate', 0):.2f}%"
                        })
                    if channel_data:
                        try:
                            st.dataframe(channel_data)
                        except ModuleNotFoundError as e:
                            # Обрабатываем ситуацию, когда не установлен pyarrow,
                            # который требуется Streamlit для отображения таблиц
                            if "pyarrow" in str(e):
                                st.warning(
                                    "Для отображения таблицы аналитики нужен пакет 'pyarrow'. "
                                    "Установите его командой: `pip install pyarrow` или "
                                    "`pip install -r requirements.txt`."
                                )
                            else:
                                raise
                
                # Рекомендации
                recommendations = st.session_state.agent.analytics.get_recommendations(report)
                if recommendations:
                    st.markdown("<h3>Рекомендации</h3>", unsafe_allow_html=True)
                    for rec in recommendations:
                        st.info(rec)
                elif report.total_posts == 0:
                    st.info("Сгенерируйте контент и создайте расписание публикаций, чтобы получить рекомендации по аналитике.")
                
                # Кнопки для скачивания отчета
                st.markdown("<h3>Скачать отчет</h3>", unsafe_allow_html=True)
                st.markdown("<p style='margin-bottom: 1rem;'>Вы можете скачать отчет по аналитике или комплексный отчет по всем разделам</p>", unsafe_allow_html=True)
                
                # Отчет по аналитике
                st.markdown("<h4>Отчет по аналитике</h4>", unsafe_allow_html=True)
                col_analytics_docx, col_analytics_pdf = st.columns(2)
                
                with col_analytics_docx:
                    try:
                        docx_bytes = generate_analytics_docx_report(
                            report, 
                            days=days,
                            product_info=st.session_state.product_info,
                            recommendations=recommendations if recommendations else []
                        )
                        st.download_button(
                            label="📄 Скачать DOCX (Аналитика)",
                            data=docx_bytes,
                            file_name=f"analytics_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Ошибка при генерации DOCX: {str(e)}")
                
                with col_analytics_pdf:
                    try:
                        pdf_bytes = generate_analytics_pdf_report(
                            report,
                            days=days,
                            product_info=st.session_state.product_info,
                            recommendations=recommendations if recommendations else []
                        )
                        st.download_button(
                            label="📄 Скачать PDF (Аналитика)",
                            data=pdf_bytes,
                            file_name=f"analytics_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Ошибка при генерации PDF: {str(e)}")
                
                # Комплексный отчет
                st.markdown("<h4>Комплексный отчет (все разделы)</h4>", unsafe_allow_html=True)
                col_comprehensive_docx, col_comprehensive_pdf = st.columns(2)
                
                with col_comprehensive_docx:
                    try:
                        comprehensive_docx = generate_comprehensive_docx_report(
                            product_info=st.session_state.get('product_info', {}),
                            deep_impact_audit=st.session_state.get('deep_impact_audit', {}),
                            hypothesis_levels=st.session_state.get('hypothesis_levels', {}),
                            strategy=st.session_state.get('strategy'),
                            content_list=st.session_state.get('content_list', []),
                            scheduling_data=st.session_state.get('scheduling_data', {}),
                            schedule=st.session_state.get('schedule', {}),
                            analytics_report=report,
                            analytics_days=days,
                            recommendations=recommendations if recommendations else []
                        )
                        st.download_button(
                            label="📄 Скачать DOCX (Все разделы)",
                            data=comprehensive_docx,
                            file_name=f"comprehensive_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Ошибка при генерации комплексного DOCX: {str(e)}")
                
                with col_comprehensive_pdf:
                    try:
                        comprehensive_pdf = generate_comprehensive_pdf_report(
                            product_info=st.session_state.get('product_info', {}),
                            deep_impact_audit=st.session_state.get('deep_impact_audit', {}),
                            hypothesis_levels=st.session_state.get('hypothesis_levels', {}),
                            strategy=st.session_state.get('strategy'),
                            content_list=st.session_state.get('content_list', []),
                            scheduling_data=st.session_state.get('scheduling_data', {}),
                            schedule=st.session_state.get('schedule', {}),
                            analytics_report=report,
                            analytics_days=days,
                            recommendations=recommendations if recommendations else []
                        )
                        st.download_button(
                            label="📄 Скачать PDF (Все разделы)",
                            data=comprehensive_pdf,
                            file_name=f"comprehensive_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Ошибка при генерации комплексного PDF: {str(e)}")
        except Exception as e:
            show_error_message(f"Ошибка при генерации отчета: {str(e)}")
            import traceback
            st.error(f"Детали ошибки: {traceback.format_exc()}")
    elif not st.session_state.content_list:
        st.info("Сначала сгенерируйте контент, чтобы увидеть аналитику.")

def main():
    """Главная функция"""
    # Инициализируем сессию перед использованием
    init_session_state()
    
    # Применяем минималистичный стиль с учетом текущей темы
    apply_minimal_style(st.session_state.theme)
    
    render_header()
    render_sidebar()
    
    # Роутинг по стадиям
    stage = st.session_state.current_stage
    
    if stage == "onboarding":
        render_onboarding()
    elif stage == "deep_impact":
        render_deep_impact()
    elif stage == "strategy":
        render_strategy()
    elif stage == "content":
        render_content()
    elif stage == "scheduling":
        render_scheduling()
    elif stage == "analytics":
        render_analytics()
    else:
        render_onboarding()

if __name__ == "__main__":
    main()

